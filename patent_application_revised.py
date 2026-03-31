# -*- coding: utf-8 -*-
"""
中国大陆发明专利申请文档生成器 - 修订版
项目名称：基于动态知识状态约束的编程教育智能体系统

修订重点：
1. 突出与对比文件CN119624716A的区别
2. 强调从零开始的学习者模型（非预设特征）
3. 强化错误概念形成-纠正机制
4. 突出知识遗忘与复习触发机制
5. 强调代码执行验证系统
6. 聚焦知识状态对LLM的动态约束
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold


def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    font_size = 16 if level == 1 else (14 if level == 2 else 12)
    set_chinese_font(run, 'SimHei', font_size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph_zh(doc, text, indent=True, bold=False, font_size=12):
    """添加中文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', font_size, bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def create_patent_application():
    """创建专利申请文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    
    # ==================== 说明书 ====================
    # 说明书标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('发明专利申请文件')
    set_chinese_font(run, 'SimHei', 22, bold=True)
    title.paragraph_format.space_after = Pt(24)
    
    # 发明名称 - 修改为更聚焦技术实现
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run('发明名称：基于动态知识状态约束的编程教育智能体系统')
    set_chinese_font(run, 'SimHei', 16, bold=True)
    name_para.paragraph_format.space_after = Pt(24)
    
    # 分隔线
    doc.add_paragraph('─' * 40)
    
    # ==================== 技术领域 ====================
    add_heading_zh(doc, '技术领域', 1)
    add_paragraph_zh(doc, 
        '本发明涉及人工智能教育技术领域，尤其涉及一种基于动态知识状态约束的编程教育智能体系统，'
        '通过实时追踪和约束大语言模型的知识掌握状态，模拟真实编程初学者的学习过程，'
        '包括知识获取、错误概念形成、知识遗忘和纠正的完整认知周期。')
    
    # ==================== 背景技术 ====================
    add_heading_zh(doc, '背景技术', 1)
    
    add_paragraph_zh(doc,
        '近年来，基于大语言模型的虚拟教学技术得到了快速发展。现有技术如CN119624716A公开了一种'
        '基于大语言模型的虚拟教学训练方法，该方法通过预设虚拟学生的先验知识水平、认知水平与学习风格'
        '等特征参数，构建具有特定能力的虚拟学生，用于教师培训和教学策略优化。')
    
    add_paragraph_zh(doc,
        '然而，现有技术存在以下技术缺陷：')
    
    add_paragraph_zh(doc,
        '1. 学习者模型静态化：现有技术采用预设特征参数（先验知识水平、认知能力、学习风格）'
        '构建虚拟学生，虚拟学生的能力边界在交互前已固定，无法模拟真实学习者从零开始的动态成长过程，'
        '导致学习过程缺乏真实性和渐进性；')
    
    add_paragraph_zh(doc,
        '2. 缺乏错误概念生命周期管理：现有技术未涉及初学者在学习过程中必然经历的错误概念形成阶段，'
        '也未提供错误概念识别、固化和纠正的完整机制，无法实现真实学习中的认知冲突与概念重构过程；')
    
    add_paragraph_zh(doc,
        '3. 知识记忆静态化处理：现有技术将知识获取视为永久性过程，未考虑艾宾浩斯遗忘曲线所描述的'
        '自然遗忘现象，无法模拟真实学习中"学习-遗忘-复习"的认知循环，也难以触发针对性的复习教学；')
    
    add_paragraph_zh(doc,
        '4. 教学反馈与代码执行脱节：在编程教育场景中，现有技术停留在文本交互层面，'
        '缺乏对学习者所教授代码的自动执行与验证能力，无法形成"教学-实践-反馈"的闭环；')
    
    add_paragraph_zh(doc,
        '5. 知识状态与语言模型响应解耦：现有技术中虚拟学生的响应基于预设角色特征，'
        '而非基于实时知识掌握状态进行动态约束，导致响应内容可能超出当前知识边界，'
        '无法实现渐进式学习的真实模拟。')
    
    add_paragraph_zh(doc,
        '因此，亟需一种能够模拟真实编程初学者完整认知过程、实现知识状态动态约束智能体响应的'
        '新型编程教育系统。')
    
    # ==================== 发明内容 ====================
    add_heading_zh(doc, '发明内容', 1)
    
    add_paragraph_zh(doc,
        '本发明的目的在于提供一种基于动态知识状态约束的编程教育智能体系统，'
        '解决现有技术中学习者模型静态化、缺乏错误概念生命周期管理、知识记忆静态化处理、'
        '教学反馈与代码执行脱节、知识状态与语言模型响应解耦的技术问题。')
    
    add_paragraph_zh(doc,
        '为实现上述目的，本发明采用如下技术方案：')
    
    add_paragraph_zh(doc,
        '一种基于动态知识状态约束的编程教育智能体系统，其特征在于，包括：')
    
    add_paragraph_zh(doc,
        '动态学习者建模模块，用于构建从零开始学习状态的AI智能体，所述AI智能体初始状态下'
        '对所有编程知识单元的掌握度为零，区别于基于预设特征参数的静态学习者模型，'
        '所述AI智能体通过教学过程动态积累知识，而非加载预设能力；')
    
    add_paragraph_zh(doc,
        '细粒度知识状态追踪模块，采用概率化知识追踪算法，用于为每个知识单元维护独立的掌握概率，'
        '所述知识状态包括：未知状态、习得状态、固化状态、遗忘状态、纠错状态五种互斥状态，'
        '实现比简单"已掌握/未掌握"二元区分更精细的学习状态刻画；')
    
    add_paragraph_zh(doc,
        '错误概念生命周期管理模块，用于模拟真实学习者的错误概念演进过程，包括：'
        '错误概念植入单元，在特定教学输入触发时，将错误理解注入AI智能体的知识表示；'
        '错误概念固化单元，通过多次强化使错误理解进入稳定状态；'
        '错误概念检测单元，基于代码执行结果和逻辑分析识别错误理解；'
        '错误概念纠正单元，通过针对性教学输入将错误状态迁移至纠错状态；')
    
    add_paragraph_zh(doc,
        '遗忘-复习触发引擎，基于艾宾浩斯遗忘曲线理论，用于动态计算每个知识单元的遗忘概率，'
        '当掌握概率因遗忘下降至阈值以下时触发复习提醒，并自适应调整遗忘速率参数，'
        '所述遗忘速率与针对该知识单元的教学交互次数负相关，实现个性化记忆强度建模；')
    
    add_paragraph_zh(doc,
        '代码执行验证与反馈模块，包括：'
        '隔离执行环境单元，用于在安全沙箱中执行AI智能体生成的编程代码；'
        '多维度评估单元，用于从语法正确性、逻辑正确性、效率指标三个维度评估代码质量；'
        '执行结果-知识状态联动单元，用于将代码执行结果反向传播至知识状态更新，'
        '正确执行提升掌握概率，错误执行标记潜在错误概念；')
    
    add_paragraph_zh(doc,
        '知识状态约束的响应生成模块，用于根据实时知识状态动态约束大语言模型的响应生成，'
        '包括：知识边界判定单元，识别当前已掌握、未掌握、存在错误理解的知识单元集合；'
        '响应约束提示构建单元，用于生成包含知识边界信息的结构化提示词；'
        '渐进响应生成单元，基于约束提示词生成仅使用已掌握知识、对未掌握知识表现困惑、'
        '对错误理解内容表现出对应错误概念的响应；')
    
    add_paragraph_zh(doc,
        '跨领域知识迁移检测模块，用于识别不同编程领域间的知识关联，当检测到已掌握领域与新领域'
        '存在共享知识单元时，触发正向迁移并加速新领域学习，实现领域间知识的动态关联与迁移。')
    
    add_paragraph_zh(doc,
        '优选的，所述细粒度知识状态追踪模块采用扩展的贝叶斯知识追踪算法，'
        '在传统BKT的习得概率、遗忘概率、猜测概率、失误概率四参数基础上，'
        '增加错误概念形成概率和错误概念纠正概率两个参数，实现六参数知识状态建模：')
    add_paragraph_zh(doc,
        'P(习得) = P(当前未掌握) × P(教学有效性) × P(前置知识满足)')
    add_paragraph_zh(doc,
        'P(遗忘) = 基础遗忘率 × 记忆强度衰减系数^(交互次数)')
    add_paragraph_zh(doc,
        'P(错误形成) = P(教学内容含误导信息) × P(学习者易错指数)')
    add_paragraph_zh(doc,
        'P(错误纠正) = P(纠错教学有效性) × P(认知冲突强度)')
    
    add_paragraph_zh(doc,
        '优选的，所述遗忘-复习触发引擎采用自适应遗忘曲线模型：')
    add_paragraph_zh(doc,
        'R(t) = e^(-t/S)')
    add_paragraph_zh(doc,
        '其中R(t)为t时刻后的记忆保持率，S为记忆强度，S = S₀ + α×N^β，'
        'S₀为基础记忆强度，N为教学交互次数，α、β为可学习的个性化参数，'
        '实现不同学习者差异化的遗忘速率建模。')
    
    add_paragraph_zh(doc,
        '优选的，所述知识状态约束的响应生成模块采用两阶段约束机制：')
    add_paragraph_zh(doc,
        '第一阶段-知识边界提取：从知识状态数据库提取当前所有知识单元的状态标记，'
        '生成结构化知识边界描述；')
    add_paragraph_zh(doc,
        '第二阶段-约束提示注入：将知识边界描述注入大语言模型的系统提示词，'
        '格式为："你已掌握：[知识列表A]；你尚未掌握：[知识列表B]；你对以下内容有错误理解：[知识列表C]"，'
        '强制模型响应严格遵循该知识边界。')
    
    add_paragraph_zh(doc,
        '本发明还提供一种基于动态知识状态约束的编程教育智能体方法，包括以下步骤：')
    
    add_paragraph_zh(doc,
        'S1：初始化阶段，构建初始知识状态矩阵，所有编程知识单元的状态设为未知，掌握概率设为0；')
    add_paragraph_zh(doc,
        'S2：教学接收阶段，接收用户输入的教学内容，解析涉及的编程知识点集合；')
    add_paragraph_zh(doc,
        'S3：知识状态更新阶段，基于概率化知识追踪算法更新各知识单元的掌握概率，'
        '若教学内容涉及前置知识且前置知识未掌握，则标记为部分习得；'
        '若教学内容存在常见错误模式，则触发错误概念形成流程；')
    add_paragraph_zh(doc,
        'S4：遗忘计算阶段，对所有已习得知识单元应用遗忘曲线模型，计算当前遗忘概率，'
        '更新掌握概率 = 原掌握概率 × (1 - 遗忘概率)，若掌握概率低于阈值则状态迁移至遗忘状态；')
    add_paragraph_zh(doc,
        'S5：约束响应生成阶段，提取当前知识边界，构建约束提示词，'
        '调用大语言模型生成受约束的学习者响应；')
    add_paragraph_zh(doc,
        'S6：代码生成阶段，根据当前知识状态选择适当难度的编程题目，触发AI智能体生成解决方案代码；')
    add_paragraph_zh(doc,
        'S7：执行验证阶段，在隔离沙箱环境中执行代码，收集语法错误、运行时错误、输出结果；')
    add_paragraph_zh(doc,
        'S8：反馈传播阶段，将执行结果反向传播至知识状态：'
        '若执行成功，提升相关知识点掌握概率；'
        '若执行失败，根据错误类型标记错误概念或降低掌握概率；')
    add_paragraph_zh(doc,
        'S9：复习触发阶段，扫描所有遗忘状态的知识单元，生成复习提醒，'
        '当用户进行复习教学时，应用更高的习得概率更新参数；')
    add_paragraph_zh(doc,
        'S10：循环执行S2-S9，直至达到预设的学习目标或循环次数。')
    
    add_paragraph_zh(doc,
        '本发明与现有技术相比的有益效果：')
    
    add_paragraph_zh(doc,
        '1. 区别于现有技术的预设特征学习者模型，本发明采用从零开始的动态学习者建模，'
        'AI智能体的知识边界在教学过程中逐步扩展，真实模拟初学者的渐进式学习过程，'
        '避免了预设能力导致的响应不真实问题；')
    
    add_paragraph_zh(doc,
        '2. 区别于现有技术的简单知识二元标记，本发明采用五状态细粒度知识追踪，'
        '特别是引入错误概念生命周期管理，完整模拟"错误形成-固化-检测-纠正"的认知过程，'
        '这是真实学习中普遍存在的现象但现有技术未涉及；')
    
    add_paragraph_zh(doc,
        '3. 区别于现有技术的静态知识记忆假设，本发明引入遗忘-复习机制，'
        '基于艾宾浩斯遗忘曲线动态计算知识遗忘，触发针对性复习，'
        '实现"学习-遗忘-复习"的认知循环模拟，增强了教学系统的真实性和持续性；')
    
    add_paragraph_zh(doc,
        '4. 区别于现有技术的纯文本交互，本发明集成代码执行验证系统，'
        '将代码执行结果与知识状态更新联动，形成"教学-代码生成-执行验证-状态更新"的完整闭环，'
        '特别适用于编程教育场景；')
    
    add_paragraph_zh(doc,
        '5. 区别于现有技术的基于预设角色的响应生成，本发明实现知识状态对LLM的动态约束，'
        '响应内容严格受限于当前知识边界，未掌握内容表现困惑，错误理解内容表现错误，'
        '实现了真正渐进式、真实的学习者行为模拟。')
    
    # ==================== 附图说明 ====================
    add_heading_zh(doc, '附图说明', 1)
    add_paragraph_zh(doc, '图1为本发明系统的整体架构示意图，展示各模块间的数据流与控制关系；')
    add_paragraph_zh(doc, '图2为细粒度知识状态追踪的状态机示意图，展示五种状态的迁移条件；')
    add_paragraph_zh(doc, '图3为错误概念生命周期管理的流程图，展示错误形成、固化、检测、纠正的完整周期；')
    add_paragraph_zh(doc, '图4为遗忘-复习触发引擎的工作原理图，展示遗忘曲线计算与复习触发机制；')
    add_paragraph_zh(doc, '图5为知识状态约束响应生成的流程图，展示知识边界提取与约束提示构建过程；')
    add_paragraph_zh(doc, '图6为代码执行验证与知识状态反馈联动的流程图；')
    add_paragraph_zh(doc, '图7为具体实施例中Python编程领域的知识单元依赖关系图示例。')
    
    # ==================== 具体实施方式 ====================
    add_heading_zh(doc, '具体实施方式', 1)
    add_paragraph_zh(doc, '下面结合具体实施例对本发明作进一步详细说明。')
    
    add_heading_zh(doc, '实施例1：系统整体架构', 2)
    add_paragraph_zh(doc,
        '如图1所示，本发明的基于动态知识状态约束的编程教育智能体系统包括以下核心组件：')
    
    add_paragraph_zh(doc,
        '前端采用React框架开发，提供学生教学界面、知识状态可视化界面、代码执行结果展示界面。'
        '学生通过聊天界面输入教学内容，系统实时展示AI智能体的学习反馈和代码执行结果。')
    
    add_paragraph_zh(doc,
        '后端采用FastAPI框架，核心服务模块包括：')
    add_paragraph_zh(doc,
        '- 动态学习者建模服务：管理AI智能体的全局状态，区别于角色预设方式，'
        '本服务的智能体初始时没有任何编程知识，所有能力通过教学交互动态构建；')
    add_paragraph_zh(doc,
        '- 细粒度知识状态追踪服务：为每个知识单元维护五状态标记和掌握概率，'
        '提供状态查询、更新、迁移功能；')
    add_paragraph_zh(doc,
        '- 错误概念生命周期服务：管理错误概念目录，检测教学内容中的错误模式，'
        '执行错误状态迁移；')
    add_paragraph_zh(doc,
        '- 遗忘-复习引擎服务：定时计算遗忘概率，触发复习提醒；')
    add_paragraph_zh(doc,
        '- 代码执行服务：在Docker容器化的隔离环境中执行Python代码，返回执行结果；')
    add_paragraph_zh(doc,
        '- 约束响应生成服务：根据知识状态构建约束提示词，调用大语言模型API生成响应；')
    
    add_paragraph_zh(doc,
        '数据层使用SQLite数据库存储知识状态矩阵、教学历史、错误概念记录、代码执行日志。')
    
    add_heading_zh(doc, '实施例2：细粒度知识状态追踪', 2)
    add_paragraph_zh(doc,
        '如图2所示，本发明采用五状态知识追踪模型，区别于传统的二元掌握标记：')
    
    add_paragraph_zh(doc,
        '未知状态（Unknown）：初始状态，掌握概率P=0，表示该知识单元从未接触。')
    add_paragraph_zh(doc,
        '习得状态（Acquired）：通过教学后掌握概率提升至阈值以上（如P≥0.7），'
        '表示已初步学会该知识单元。')
    add_paragraph_zh(doc,
        '固化状态（Solidified）：通过多次成功应用和测试，掌握概率进一步提升（如P≥0.9），'
        '遗忘速率降低，表示知识已进入长期记忆。')
    add_paragraph_zh(doc,
        '遗忘状态（Forgotten）：因长时间未复习，掌握概率下降至阈值以下，'
        '需要触发复习教学。')
    add_paragraph_zh(doc,
        '纠错状态（Corrected）：曾经存在错误理解，经纠错教学后纠正，'
        '掌握概率从零开始重新积累。')
    
    add_paragraph_zh(doc,
        '状态迁移条件如下：')
    add_paragraph_zh(doc,
        '未知→习得：教学输入覆盖该知识单元，且前置知识满足，习得概率计算结果达到阈值。')
    add_paragraph_zh(doc,
        '习得→固化：多次（如3次以上）成功应用该知识单元解决问题，掌握概率累积提升。')
    add_paragraph_zh(doc,
        '习得/固化→遗忘：经过时间t，R(t) = e^(-t/S) < 阈值，掌握概率相应衰减。')
    add_paragraph_zh(doc,
        '未知/习得→纠错：教学内容触发错误概念形成，或代码执行错误被识别为特定错误模式。')
    add_paragraph_zh(doc,
        '纠错→习得：纠错教学成功，认知冲突解决，重新习得。')
    
    add_heading_zh(doc, '实施例3：错误概念生命周期管理', 2)
    add_paragraph_zh(doc,
        '如图3所示，本发明完整模拟真实学习者的错误概念演进过程：')
    
    add_paragraph_zh(doc,
        '错误概念植入：系统维护编程领域常见错误概念目录，如Python中的混淆赋值与相等运算符、循环变量作用域误解等。当检测教学内容包含错误信息，'
        '或代码执行产生特定错误模式时，将对应错误概念标记植入知识状态。')
    
    add_paragraph_zh(doc,
        '错误概念固化：若错误概念植入后未得到及时纠正，且后续教学或代码执行中多次强化该错误模式，'
        '错误概念进入固化状态，掌握概率在错误理解维度累积。')
    
    add_paragraph_zh(doc,
        '错误概念检测：通过两种途径检测：'
        '（1）静态分析：解析AI智能体生成的代码，匹配已知错误模式；'
        '（2）动态执行：在沙箱中执行代码，捕获运行时错误，映射至错误概念目录。')
    
    add_paragraph_zh(doc,
        '错误概念纠正：当检测到错误概念后，系统调整AI智能体的响应策略，'
        '使其表现出对该知识点的困惑或错误理解，引导学生进行纠错教学。'
        '纠错教学内容经解析后，触发错误概念状态迁移至纠错状态，'
        '并应用错误纠正概率模型重新计算掌握概率。')
    
    add_heading_zh(doc, '实施例4：遗忘-复习触发引擎', 2)
    add_paragraph_zh(doc,
        '如图4所示，遗忘-复习引擎基于艾宾浩斯遗忘曲线理论，但做了编程教育场景的自适应改进：')
    
    add_paragraph_zh(doc,
        '基础遗忘模型：R(t) = e^(-t/S)，其中t为距上次教学交互的时间（天），'
        'S为记忆强度。')
    
    add_paragraph_zh(doc,
        '自适应记忆强度：S = S₀ + α×N^β，其中S₀=1（基础记忆强度），'
        'N为该知识单元的教学交互次数，α=0.3，β=0.5为通过实验数据拟合的参数。'
        '这意味着随着复习次数增加，记忆强度提升，遗忘速率降低。')
    
    add_paragraph_zh(doc,
        '遗忘触发与复习：系统每日扫描所有已习得知识单元，计算当前R(t)。'
        '若R(t) < 0.5（即掌握概率衰减至50%以下），触发复习提醒。'
        '复习教学内容应用更高的习得概率更新（如P(习得)=0.9而非0.7），'
        '模拟复习的高效性。')
    
    add_heading_zh(doc, '实施例5：知识状态约束的响应生成', 2)
    add_paragraph_zh(doc,
        '如图5所示，区别于现有技术基于预设角色的响应生成，本发明实现动态知识状态约束：')
    
    add_paragraph_zh(doc,
        '知识边界提取：在每次响应生成前，系统从知识状态数据库查询所有知识单元的状态，'
        '分类为：')
    add_paragraph_zh(doc, '- 已完全掌握：掌握概率≥0.9')
    add_paragraph_zh(doc, '- 已初步掌握：0.7≤掌握概率<0.9')
    add_paragraph_zh(doc, '- 未掌握但前置知识满足：可以学习')
    add_paragraph_zh(doc, '- 未掌握且前置知识不满足：无法理解')
    add_paragraph_zh(doc, '- 存在错误理解：错误概念状态标记')
    
    add_paragraph_zh(doc,
        '约束提示构建：将上述分类构建为结构化提示词，例如：')
    add_paragraph_zh(doc,
        '"作为编程初学者，你的知识状态如下：\n'
        '- 你已熟练掌握：变量定义、基本运算\n'
        '- 你初步掌握但还不熟练：if条件判断\n'
        '- 你尚未学习：for循环、函数定义\n'
        '- 你对"==与=的区别"存在错误理解，认为两者相同\n'
        '请基于以上知识状态，以第一人称回应用户的教学。'
        '只能使用已掌握的知识，对未掌握内容表现困惑和好奇，'
        '对存在错误理解的内容表现出对应的错误。"')
    
    add_paragraph_zh(doc,
        '渐进响应生成：将约束提示词与用户的教学输入一并发送至大语言模型API，'
        '生成的响应严格遵循知识边界，实现渐进式、真实的学习者行为模拟。')
    
    add_heading_zh(doc, '实施例6：代码执行验证与反馈联动', 2)
    add_paragraph_zh(doc,
        '如图6所示，代码执行模块与知识状态追踪形成闭环：')
    
    add_paragraph_zh(doc,
        '隔离执行环境：使用Docker容器创建隔离沙箱，限制执行时间（如5秒）、内存（如128MB），'
        '防止恶意代码。')
    
    add_paragraph_zh(doc,
        '多维度评估：')
    add_paragraph_zh(doc, '- 语法维度：编译/解析是否通过')
    add_paragraph_zh(doc, '- 逻辑维度：输出结果是否符合预期')
    add_paragraph_zh(doc, '- 效率维度：时间复杂度、空间复杂度评估')
    
    add_paragraph_zh(doc,
        '反馈传播机制：')
    add_paragraph_zh(doc,
        '- 执行成功：涉及的所有知识单元掌握概率提升10%，若已存在错误概念标记则触发纠错流程')
    add_paragraph_zh(doc,
        '- 语法错误：若错误匹配已知错误概念目录（如"缩进错误"对应"Python语法结构误解"），'
        '则标记该错误概念')
    add_paragraph_zh(doc,
        '- 逻辑错误：分析错误类型，若为算法逻辑错误则降低对应算法知识单元掌握概率，'
        '若为概念误解则标记概念错误')
    
    add_heading_zh(doc, '实施例7：跨领域知识迁移检测', 2)
    add_paragraph_zh(doc,
        '如图7所示，系统支持Python编程、数据库SQL、AI素养等多个领域，'
        '各领域独立定义知识单元，但系统检测跨领域知识关联：')
    
    add_paragraph_zh(doc,
        '当用户在Python领域已掌握"条件判断"知识单元后，开始学习数据库领域的"WHERE条件筛选"时，'
        '系统检测到两个知识单元在逻辑结构上的相似性（都是条件筛选），'
        '触发正向迁移：数据库领域的"WHERE条件筛选"初始掌握概率设为0.3而非0，'
        '遗忘速率降低，实现跨领域学习效率提升。')
    
    add_heading_zh(doc, '实施例8：完整教学流程示例', 2)
    add_paragraph_zh(doc,
        '以Python编程初学者教授AI智能体"for循环"概念为例：')
    
    add_paragraph_zh(doc,
        '初始状态：AI智能体对for循环完全未知，掌握概率为0，但已掌握变量定义和print函数。')
    
    add_paragraph_zh(doc,
        '第一次教学：用户解释for循环的基本语法。系统更新知识状态：for循环标记为习得，P=0.7。'
        'AI智能体根据约束提示生成响应：表现出理解但还不熟练，尝试用for循环写简单代码。')
    
    add_paragraph_zh(doc,
        '代码生成与执行：AI生成一个遍历列表的代码。代码在沙箱中执行成功，'
        'for循环掌握概率提升至0.8。')
    
    add_paragraph_zh(doc,
        '错误概念形成：用户无意中说出"for循环和while循环完全一样"。'
        '系统检测到错误信息，标记"循环类型混淆"错误概念。')
    
    add_paragraph_zh(doc,
        '错误概念表现：后续对话中，当讨论循环选择时，AI表现出对两种循环适用场景的错误理解。')
    
    add_paragraph_zh(doc,
        '错误纠正：用户纠正上述错误，解释for和while的区别。系统更新错误概念状态为纠错状态，'
        '重新计算相关掌握概率。')
    
    add_paragraph_zh(doc,
        '遗忘与复习：经过7天未接触for循环，根据遗忘曲线R(7)=e^(-7/2.3)≈0.48，'
        '掌握概率从0.8衰减至0.38，触发复习提醒。用户进行复习教学，'
        '应用高习得概率更新后，掌握概率恢复至0.85。')
    
    # 分页
    doc.add_page_break()
    
    # ==================== 权利要求书 ====================
    rights_title = doc.add_paragraph()
    rights_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rights_title.add_run('权利要求书')
    set_chinese_font(run, 'SimHei', 18, bold=True)
    rights_title.paragraph_format.space_after = Pt(18)
    
    # 权利要求1 - 独立权利要求（系统）
    p = doc.add_paragraph()
    run = p.add_run('1. 一种基于动态知识状态约束的编程教育智能体系统，其特征在于，包括：')
    set_chinese_font(run, 'SimSun', 12)
    
    add_paragraph_zh(doc, 
        '动态学习者建模模块，用于构建从零开始学习状态的AI智能体，所述AI智能体初始状态下对所有编程知识单元的掌握度为零，通过教学过程动态积累知识而非加载预设能力；', 
        indent=True)
    add_paragraph_zh(doc, 
        '细粒度知识状态追踪模块，采用概率化知识追踪算法，用于为每个知识单元维护独立的掌握概率，所述知识状态包括未知状态、习得状态、固化状态、遗忘状态、纠错状态五种互斥状态；', 
        indent=True)
    add_paragraph_zh(doc, 
        '错误概念生命周期管理模块，用于模拟真实学习者的错误概念演进过程，包括错误概念植入单元、错误概念固化单元、错误概念检测单元和错误概念纠正单元；', 
        indent=True)
    add_paragraph_zh(doc, 
        '遗忘-复习触发引擎，基于艾宾浩斯遗忘曲线理论，用于动态计算每个知识单元的遗忘概率，当掌握概率因遗忘下降至阈值以下时触发复习提醒；', 
        indent=True)
    add_paragraph_zh(doc, 
        '代码执行验证与反馈模块，包括隔离执行环境单元、多维度评估单元和执行结果-知识状态联动单元，用于将代码执行结果反向传播至知识状态更新；', 
        indent=True)
    add_paragraph_zh(doc, 
        '知识状态约束的响应生成模块，用于根据实时知识状态动态约束大语言模型的响应生成，包括知识边界判定单元、响应约束提示构建单元和渐进响应生成单元。', 
        indent=True)
    
    # 权利要求2
    add_paragraph_zh(doc, 
        '2. 根据权利要求1所述的系统，其特征在于，所述细粒度知识状态追踪模块采用扩展的贝叶斯知识追踪算法，在传统四参数基础上增加错误概念形成概率和错误概念纠正概率两个参数，实现六参数知识状态建模。',
        indent=False)
    
    # 权利要求3
    add_paragraph_zh(doc, 
        '3. 根据权利要求1所述的系统，其特征在于，所述错误概念生命周期管理模块的错误概念植入单元用于在检测教学内容含误导信息或代码执行产生特定错误模式时，将对应错误概念标记植入知识状态。',
        indent=False)
    
    # 权利要求4
    add_paragraph_zh(doc, 
        '4. 根据权利要求1所述的系统，其特征在于，所述遗忘-复习触发引擎采用自适应遗忘曲线模型R(t) = e^(-t/S)，其中记忆强度S = S₀ + α×N^β，N为教学交互次数，α、β为个性化参数，实现不同知识单元的差异化遗忘速率建模。',
        indent=False)
    
    # 权利要求5
    add_paragraph_zh(doc, 
        '5. 根据权利要求1所述的系统，其特征在于，所述代码执行验证与反馈模块的隔离执行环境单元使用容器化技术创建隔离沙箱，限制执行时间和内存资源；所述多维度评估单元从语法正确性、逻辑正确性、效率指标三个维度评估代码质量。',
        indent=False)
    
    # 权利要求6
    add_paragraph_zh(doc, 
        '6. 根据权利要求1所述的系统，其特征在于，所述知识状态约束的响应生成模块采用两阶段约束机制：第一阶段提取当前知识边界，第二阶段将知识边界描述注入大语言模型的系统提示词，强制模型响应严格遵循该知识边界。',
        indent=False)
    
    # 权利要求7
    add_paragraph_zh(doc, 
        '7. 根据权利要求6所述的系统，其特征在于，所述知识边界描述格式为："你已掌握：[知识列表A]；你尚未掌握：[知识列表B]；你对以下内容有错误理解：[知识列表C]"，所述渐进响应生成单元基于该约束提示词生成仅使用已掌握知识、对未掌握知识表现困惑、对错误理解内容表现出对应错误概念的响应。',
        indent=False)
    
    # 权利要求8
    add_paragraph_zh(doc, 
        '8. 根据权利要求1所述的系统，其特征在于，还包括跨领域知识迁移检测模块，用于识别不同编程领域间的知识关联，当检测到已掌握领域与新领域存在共享知识单元时，触发正向迁移并加速新领域学习。',
        indent=False)
    
    # 权利要求9 - 独立权利要求（方法）
    add_paragraph_zh(doc, 
        '9. 一种基于动态知识状态约束的编程教育智能体方法，其特征在于，包括以下步骤：',
        indent=False)
    add_paragraph_zh(doc, 'S1：初始化阶段，构建初始知识状态矩阵，所有编程知识单元的状态设为未知，掌握概率设为0；', indent=True)
    add_paragraph_zh(doc, 'S2：教学接收阶段，接收用户输入的教学内容，解析涉及的编程知识点集合；', indent=True)
    add_paragraph_zh(doc, 'S3：知识状态更新阶段，基于概率化知识追踪算法更新各知识单元的掌握概率，若教学内容存在常见错误模式则触发错误概念形成流程；', indent=True)
    add_paragraph_zh(doc, 'S4：遗忘计算阶段，对所有已习得知识单元应用遗忘曲线模型，计算当前遗忘概率并更新掌握概率，若低于阈值则状态迁移至遗忘状态；', indent=True)
    add_paragraph_zh(doc, 'S5：约束响应生成阶段，提取当前知识边界，构建约束提示词，调用大语言模型生成受约束的学习者响应；', indent=True)
    add_paragraph_zh(doc, 'S6：代码生成阶段，根据当前知识状态选择适当难度的编程题目，触发AI智能体生成解决方案代码；', indent=True)
    add_paragraph_zh(doc, 'S7：执行验证阶段，在隔离沙箱环境中执行代码，收集语法错误、运行时错误、输出结果；', indent=True)
    add_paragraph_zh(doc, 'S8：反馈传播阶段，将执行结果反向传播至知识状态：若执行成功则提升掌握概率，若执行失败则根据错误类型标记错误概念或降低掌握概率；', indent=True)
    add_paragraph_zh(doc, 'S9：复习触发阶段，扫描所有遗忘状态的知识单元，生成复习提醒，当用户进行复习教学时应用更高的习得概率更新参数；', indent=True)
    add_paragraph_zh(doc, 'S10：循环执行S2-S9，直至达到预设的学习目标。', indent=True)
    
    # 权利要求10
    add_paragraph_zh(doc, 
        '10. 根据权利要求9所述的方法，其特征在于，步骤S3中采用六参数知识状态建模，所述六参数包括习得概率、遗忘概率、猜测概率、失误概率、错误概念形成概率和错误概念纠正概率。',
        indent=False)
    
    # 分页
    doc.add_page_break()
    
    # ==================== 摘要 ====================
    abstract_title = doc.add_paragraph()
    abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = abstract_title.add_run('摘  要')
    set_chinese_font(run, 'SimHei', 18, bold=True)
    abstract_title.paragraph_format.space_after = Pt(18)
    
    abstract_text = (
        '本发明公开了一种基于动态知识状态约束的编程教育智能体系统及方法，区别于现有技术基于预设特征参数的学习者模型，'
        '本发明构建从零开始学习状态的AI智能体，通过细粒度五状态知识追踪、错误概念生命周期管理、遗忘-复习触发引擎、'
        '代码执行验证与反馈联动、知识状态约束的响应生成等技术手段，实现真实编程初学者完整认知过程的动态模拟。'
        '系统为每个知识单元维护未知、习得、固化、遗忘、纠错五种互斥状态，模拟错误概念的形成-固化-检测-纠正周期，'
        '基于艾宾浩斯遗忘曲线触发个性化复习，并将代码执行结果反向传播至知识状态更新，最终通过知识边界约束实现大语言模型响应的渐进式生成。'
    )
    add_paragraph_zh(doc, abstract_text, indent=True)
    
    # 关键词
    keywords = doc.add_paragraph()
    run = keywords.add_run('关键词：')
    set_chinese_font(run, 'SimHei', 12, bold=True)
    run = keywords.add_run('动态知识状态约束；从零开始学习者模型；错误概念生命周期；遗忘-复习触发；代码执行反馈；渐进式响应生成')
    set_chinese_font(run, 'SimSun', 12)
    keywords.paragraph_format.first_line_indent = Cm(0.74)
    
    # 保存文档
    output_path = 'e:\\cs teachable agent\\发明专利申请文件-基于动态知识状态约束的编程教育智能体系统-修订版.docx'
    doc.save(output_path)
    print(f'专利申请修订版文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    create_patent_application()
