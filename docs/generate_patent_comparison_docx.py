# -*- coding: utf-8 -*-
"""生成《专利对比与创造性分析》Word 文档，与 Markdown 内容一致。"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, name="SimSun", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, "SimHei", 16, True)
    doc.add_paragraph()


def add_h(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, "SimHei", 14 if level == 1 else 12, True)


def add_p(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r, "SimSun", 11)
    p.paragraph_format.line_spacing = 1.5


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                set_run_font(run, "SimHei", 10, True)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, cell in enumerate(row):
            cells[ci].text = str(cell)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    set_run_font(run, "SimSun", 9)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    add_title(doc, "与已公开对比文件的区别点及创造性初步分析")

    add_p(
        doc,
        "文档用途：供专利代理机构撰写意见陈述书、新颖性/创造性答复或申请策略参考。",
        indent=False,
    )
    add_p(
        doc,
        "对比文件：CN119624716A《基于大语言模型的虚拟教学训练方法、系统及介质》（以国家知识产权局公开文本为准）。",
        indent=False,
    )
    add_p(
        doc,
        "我方方案：编程教育场景下可教智能体系统——学习者向大模型扮演的虚拟学习者施教，以动态知识状态驱动行为，并与代码结构校验、沙箱执行、任务调度、过程追溯等形成闭环。",
        indent=False,
    )

    add_h(doc, "一、对比文件（CN119624716A）技术要点摘要", 1)
    add_table(
        doc,
        ["维度", "对比文件公开内容（概括）"],
        [
            (
                "应用场景",
                "面向教师/培训师的虚拟教学训练；提升教学质量与策略。",
            ),
            (
                "虚拟学生构建",
                "按教学层次与学科，预设先验知识水平、认知水平、学习风格等特征参数，由大模型生成虚拟学生。",
            ),
            (
                "交互形态",
                "多模态教学信息输入；虚拟学生学习并产生学习状态与反馈。",
            ),
            (
                "评估与建议",
                "虚拟教学专家系统评估教学过程，输出策略优化建议并下发终端。",
            ),
            (
                "硬件",
                "计算机、麦克风、摄像头等采集多模态信息。",
            ),
        ],
    )
    doc.add_paragraph()

    add_h(doc, "二、我方方案技术要点摘要", 1)
    add_table(
        doc,
        ["维度", "我方要点"],
        [
            ("应用场景", "面向学生编程学习；以教促学；可含教师分析端，核心闭环在学-教-测-状态。"),
            ("虚拟学习者起点", "非预设能力档位；以知识单元为粒度从零或极低先验开始动态演化。"),
            ("知识建模", "扩展状态（含错误概念、证据等）；BKT思路+依赖门控+教学质量等因素。"),
            ("行为约束", "知识状态约束的大模型输出（如Reflect-Respond管道）。"),
            ("编程特有", "AST节点类型映射知识单元；约束可生成/可执行代码与知识边界一致。"),
            ("执行闭环", "沙箱执行+错误归因；结果回写知识状态与错误概念。"),
            ("调度", "多策略任务选择；SM-2类间隔与掌握度联动（若实现）。"),
            ("可追溯", "多类型事件记录教学-状态-对话-选题-尝试-评估的因果关联。"),
        ],
    )
    doc.add_paragraph()

    add_h(doc, "三、区别点对照表（代理可直接引用）", 1)
    add_table(
        doc,
        ["序号", "区别主题", "对比文件", "我方", "是否易区分"],
        [
            ("D1", "应用目的与用户角色", "培训教师；专家给教学策略建议", "学生教虚拟学习者；核心闭环在自身学习效果", "较易"),
            ("D2", "虚拟学习者能力来源", "预设先验/认知/学习风格等参数", "知识单元+动态掌握概率/状态演化", "较易"),
            ("D3", "多模态vs编程闭环", "文本/图/视频等通用材料", "源码+沙箱+测评与知识状态联动", "较易"),
            ("D4", "专家评估模块", "独立虚拟专家评估并建议", "主路径为状态机+约束生成+执行反馈", "较易"),
            ("D5", "知识粒度与状态", "学习状态较宏观", "细粒度单元、前置依赖、多类证据", "中等"),
            ("D6", "错误概念", "未公开同构方案", "目录+触发+严重度/衰减+与测评联动", "中等偏强"),
            ("D7", "遗忘与复习", "未公开同构融合", "个性化衰减+与掌握度联动调间隔", "中等"),
            ("D8", "AST语法守卫", "未涉及", "语法节点→知识单元映射与子集校验", "较强"),
            ("D9", "多策略选题", "未公开同构", "按错误概念/不确定性/间隔等切换", "中等偏强"),
            ("D10", "Reflect-Respond", "未公开同构管道", "分阶段+严格仅用已检索知识的提示", "中等"),
            ("D11", "证据链", "未公开同构", "多事件类型+快照+关联ID", "中等"),
            ("D12", "模式切换提问", "未公开", "周期性苏格拉底式追问", "偏弱至中等"),
        ],
    )
    doc.add_paragraph()

    add_h(doc, "四、各区别点创造性初步判断（非法律结论）", 1)
    add_p(
        doc,
        "说明：下列为技术交底层面粗判，最终以审查员检索及代理法律分析为准。",
        indent=False,
    )
    add_table(
        doc,
        ["区别点", "新颖性(相对716A)", "创造性粗判", "撰写/举证侧重"],
        [
            ("D1", "通常有区别", "中等；需绑定D5-D11", "作场景限定，不单作主创新"),
            ("D2", "有区别", "中等偏强", "状态来源：教学/测评/执行结果"),
            ("D3", "有区别", "较强", "执行结果→知识单元/错误概念的映射规则"),
            ("D4", "有区别", "弱至中等", "区别说明即可，不单列核心权要"),
            ("D5", "部分有区别", "偏弱", "与D8/D3/D7协同或非显而易见参数化"),
            ("D6", "有区别", "中等偏强", "触发源、阈值、与掌握度联动"),
            ("D7", "有区别", "中等", "融合规则/公式，非仅陈述遗忘曲线"),
            ("D8", "有区别", "较强", "保留在重要权利要求；可附映射表"),
            ("D9", "有区别", "中等", "输入变量与优先级/切换条件具体化"),
            ("D10", "表述有区别", "偏弱至中等", "绑定状态库+AST+执行反馈"),
            ("D11", "有区别", "偏弱", "与每次状态转移因果绑定"),
            ("D12", "有区别", "偏弱", "从属或说明书效果"),
        ],
    )
    doc.add_paragraph()

    add_h(doc, "五、审查风险与应对要点", 1)
    add_p(
        doc,
        "1）以教促学+LLM虚拟学生：可能被其他文献覆盖——应对：锚定D3、D8、D6、D7联动与可验证实现。",
        indent=False,
    )
    add_p(
        doc,
        "2）BKT、遗忘曲线、沙箱各自公知：易被指常规组合——应对：写清组合解决的技术问题（防全知、代码与状态一致、复习间隔随掌握度）。",
        indent=False,
    )
    add_p(
        doc,
        "3）提示词/管道创造性不足——应对：与状态矩阵、AST允许集合、沙箱反馈写入同一权利要求链或方法步骤。",
        indent=False,
    )
    add_p(
        doc,
        "4）对比文件「学习状态信息」——应对：用D5/D6/D8的数据结构、状态种类、更新触发源细化差异。",
        indent=False,
    )

    add_h(doc, "六、一句话结论（非法律结论）", 1)
    add_p(
        doc,
        "与716A最易拉开距离：编程场景下AST-知识单元绑定+沙箱执行驱动知识状态/错误概念（D3、D8），以及动态知识边界约束与测评/选题闭环（D2、D5、D9组合）；单独「逆向教学」叙事支撑弱。",
        indent=False,
    )
    add_p(
        doc,
        "单独支撑创造性偏弱：泛泛BKT、沙箱、遗忘曲线、日志、纯提示工程——需用联动与具体规则补强。",
        indent=False,
    )
    add_p(
        doc,
        "是否具备创造性须充分检索后定论；从交底角度，D3+D8+D6+D7协同较适合作为答复重点。",
        indent=False,
    )

    out = Path(__file__).resolve().parent / "专利对比与创造性分析-与CN119624716A区别点.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
