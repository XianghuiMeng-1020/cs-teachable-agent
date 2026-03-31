# -*- coding: utf-8 -*-
"""Generate Patent Comparison Analysis Document (English filename to avoid encoding issues)"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, name="SimSun", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading_zh(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font_size = 16 if level == 1 else (14 if level == 2 else 12)
    set_run_font(run, 'SimHei', font_size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph_zh(doc, text, indent=True, bold=False, font_size=11):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, 'SimSun', font_size, bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("与CN119624716A的深度区别分析及创造性评估")
    set_run_font(run, "SimHei", 18, True)
    doc.add_paragraph()

    # Meta info
    add_paragraph_zh(doc, "文档用途：供专利代理机构撰写意见陈述书、新颖性/创造性答复使用", indent=False)
    add_paragraph_zh(doc, "对比文件：CN119624716A《基于大语言模型的虚拟教学训练方法、系统及介质》", indent=False)
    add_paragraph_zh(doc, "生成日期：2026年3月20日", indent=False)
    doc.add_paragraph()

    # Section 1: 对方专利核心技术分析
    add_heading_zh(doc, "一、对比文件（CN119624716A）核心技术解构", 1)
    
    add_heading_zh(doc, "1.1 技术架构", 2)
    add_paragraph_zh(doc, "对比文件采用三层架构：")
    add_paragraph_zh(doc, "1) 虚拟学生模块：基于预设参数（先验知识水平、认知水平、学习风格）构建虚拟学生")
    add_paragraph_zh(doc, "2) 虚拟专家模块：阶段性评估教学效果，提出改进建议")
    add_paragraph_zh(doc, "3) 硬件设备：采集文本/语音/图像进行多模态交互")
    
    add_heading_zh(doc, "1.2 关键技术特征", 2)
    add_paragraph_zh(doc, "【特征A】静态能力预设：在交互前已固定虚拟学生的能力边界")
    add_paragraph_zh(doc, "【特征B】专家评估驱动：依赖独立专家系统输出策略建议")
    add_paragraph_zh(doc, "【特征C】多模态输入：支持文本、图像、视频等教学材料")
    add_paragraph_zh(doc, "【特征D】教师培训导向：核心目的是提升教师教学策略")
    
    add_heading_zh(doc, "1.3 技术缺陷（本方案解决的目标）", 2)
    add_paragraph_zh(doc, "1) 虚拟学生能力静态化：无法模拟从零开始的动态成长过程")
    add_paragraph_zh(doc, "2) 缺乏知识状态精确追踪：仅有宏观学习状态，无细粒度知识单元管理")
    add_paragraph_zh(doc, "3) 无错误概念机制：未涉及初学者必然经历的错误理解形成与纠正")
    add_paragraph_zh(doc, "4) 无代码执行验证：停留在文本交互层面，无法形成教学-实践闭环")
    add_paragraph_zh(doc, "5) 响应生成与知识状态解耦：基于预设角色而非实时知识边界")
    doc.add_paragraph()

    # Section 2: 我方专利核心技术分析
    add_heading_zh(doc, "二、我方方案核心技术解构", 1)
    
    add_heading_zh(doc, "2.1 技术架构（8大核心模块）", 2)
    add_paragraph_zh(doc, "模块1：六参数贝叶斯知识追踪（6-Param BKT）")
    add_paragraph_zh(doc, "模块2：抽象语法树级代码守卫（AST Guard）")
    add_paragraph_zh(doc, "模块3：多策略自适应任务选择")
    add_paragraph_zh(doc, "模块4：完整证据链追踪系统")
    add_paragraph_zh(doc, "模块5：SM-2与BKT融合调度")
    add_paragraph_zh(doc, "模块6：Reflect-Respond四阶段约束响应")
    add_paragraph_zh(doc, "模块7：智能模式切换（苏格拉底提问）")
    add_paragraph_zh(doc, "模块8：教学解释多维质量评估")
    
    add_heading_zh(doc, "2.2 关键技术特征", 2)
    add_paragraph_zh(doc, "【特征A】动态能力构建：从零开始，通过教学交互动态积累知识")
    add_paragraph_zh(doc, "【特征B】六参数知识追踪：习得/遗忘/猜测/失误/错误形成/错误纠正")
    add_paragraph_zh(doc, "【特征C】AST级代码守卫：语法结构节点精确映射到知识单元")
    add_paragraph_zh(doc, "【特征D】执行反馈闭环：沙箱执行结果反向驱动知识状态更新")
    add_paragraph_zh(doc, "【特征E】知识状态约束生成：实时知识边界通过提示词强制约束LLM响应")
    add_paragraph_zh(doc, "【特征F】学生导向：以教促学，侧重学习者知识内化")
    doc.add_paragraph()

    # Section 3: 逐条深度对比分析
    add_heading_zh(doc, "三、深度区别分析（12个维度逐条对比）", 1)
    
    # D1
    add_heading_zh(doc, "D1: 应用场景与用户角色", 2)
    add_paragraph_zh(doc, "【对方】教师培训场景，虚拟专家系统给教师提供教学策略优化建议")
    add_paragraph_zh(doc, "【我方】学生自学场景，学生向虚拟学习者施教实现以教促学")
    add_paragraph_zh(doc, "【区别本质】对方是「专家辅助系统」，我方是「学习闭环系统」")
    add_paragraph_zh(doc, "【创造性判断】★★☆☆☆ 单独应用场景差异创造性较弱，需结合技术实现差异")
    
    # D2
    add_heading_zh(doc, "D2: 虚拟学习者能力来源（核心区别）", 2)
    add_paragraph_zh(doc, "【对方】预设参数模型：")
    add_paragraph_zh(doc, "  - 先验知识水平：高中/本科/研究生等固定档位")
    add_paragraph_zh(doc, "  - 认知水平：基础/中等/高级三档")
    add_paragraph_zh(doc, "  - 学习风格：视觉/听觉/动觉三种类型")
    add_paragraph_zh(doc, "  → 能力边界在交互前已确定，交互中保持不变")
    add_paragraph_zh(doc, "【我方】动态演化模型：")
    add_paragraph_zh(doc, "  - 初始状态：所有知识单元掌握概率为0.01（接近零起点）")
    add_paragraph_zh(doc, "  - 演化机制：通过教学输入动态更新掌握概率")
    add_paragraph_zh(doc, "  - 状态种类：未知/部分习得/已习得/已固化/已遗忘/错误概念/已纠正")
    add_paragraph_zh(doc, "  → 能力边界随教学过程动态扩展")
    add_paragraph_zh(doc, "【区别本质】静态预设 vs 动态演化；粗粒度档位 vs 细粒度单元")
    add_paragraph_zh(doc, "【创造性判断】★★★★☆ 动态演化模型与静态预设形成显著对立，具有较强创造性")
    
    # D3
    add_heading_zh(doc, "D3: 知识追踪模型（核心区别）", 2)
    add_paragraph_zh(doc, "【对方】未公开细粒度知识追踪，仅提及「学习状态信息」")
    add_paragraph_zh(doc, "【我方】六参数贝叶斯知识追踪：")
    add_paragraph_zh(doc, "  参数1：习得概率 P(T) = 0.1（从未掌握到掌握的基线概率）")
    add_paragraph_zh(doc, "  参数2：遗忘概率 P(F) = 0.1（基础遗忘率，实际计算中与交互次数负相关）")
    add_paragraph_zh(doc, "  参数3：猜测概率 P(G) = 0.25（未掌握却蒙对的概率）")
    add_paragraph_zh(doc, "  参数4：失误概率 P(S) = 0.1（掌握后仍犯错的概率）")
    add_paragraph_zh(doc, "  参数5：错误概念形成概率 P(M) = 0.3（教学内容含误导时触发）")
    add_paragraph_zh(doc, "  参数6：错误概念纠正概率 P(C) = 0.5（纠错教学有效性决定）")
    add_paragraph_zh(doc, "【公式对比】")
    add_paragraph_zh(doc, "  传统BKT（四参数）：P(L_new) = P(L) + (1-P(L)) × P(T)")
    add_paragraph_zh(doc, "  六参数BKT：P(L_new) = P(L) + (1-P(L)) × P(T) × QualityScore × PrereqSatisfied")
    add_paragraph_zh(doc, "【区别本质】四参数基础模型 vs 六参数扩展模型；新增错误概念生命周期参数")
    add_paragraph_zh(doc, "【创造性判断】★★★★☆ 六参数扩展具有明确的认知心理学依据，参数公式具体化，创造性较强")
    
    # D4
    add_heading_zh(doc, "D4: 错误概念管理机制（核心区别）", 2)
    add_paragraph_zh(doc, "【对方】完全未涉及错误概念相关内容")
    add_paragraph_zh(doc, "【我方】完整错误概念生命周期管理：")
    add_paragraph_zh(doc, "  阶段1-错误概念植入：")
    add_paragraph_zh(doc, "    - 触发条件1：教学内容包含误导信息（ContentMisleading=1）")
    add_paragraph_zh(doc, "    - 触发条件2：代码执行失败匹配已知错误模式")
    add_paragraph_zh(doc, "    - 触发条件3：LLM诊断置信度>0.6时自动标记")
    add_paragraph_zh(doc, "    - 植入公式：P(M_triggered) = ContentMisleading × LearnerErrorProneIndex")
    add_paragraph_zh(doc, "  阶段2-错误概念固化：")
    add_paragraph_zh(doc, "    - 多次强化后进入稳定状态，严重度评分累积")
    add_paragraph_zh(doc, "  阶段3-错误概念检测：")
    add_paragraph_zh(doc, "    - 静态分析：AST解析匹配错误模式（如混淆=与==）")
    add_paragraph_zh(doc, "    - 动态执行：沙箱执行捕获运行时错误")
    add_paragraph_zh(doc, "  阶段4-错误概念纠正：")
    add_paragraph_zh(doc, "    - 触发：检测到纠错教学内容")
    add_paragraph_zh(doc, "    - 公式：P(C) = CorrectionQuality × CognitiveConflictIntensity")
    add_paragraph_zh(doc, "    - 效果：状态迁移至纠错，掌握概率重置重新积累")
    add_paragraph_zh(doc, "  阶段5-错误概念衰减：")
    add_paragraph_zh(doc, "    - 每次正确教学后严重度×0.85，低于0.2阈值时移除")
    add_paragraph_zh(doc, "【区别本质】无错误概念管理 vs 五阶段完整生命周期管理")
    add_paragraph_zh(doc, "【创造性判断】★★★★★ 错误概念形成-固化-检测-纠正-衰减的完整机制具有显著创造性，教育心理学理论支撑明确")
    
    # D5
    add_heading_zh(doc, "D5: 遗忘与复习机制（核心区别）", 2)
    add_paragraph_zh(doc, "【对方】未涉及遗忘机制，假设知识一旦获取永久保持")
    add_paragraph_zh(doc, "【我方】艾宾浩斯遗忘曲线+个性化调整：")
    add_paragraph_zh(doc, "  基础模型：R(t) = e^(-t/S)")
    add_paragraph_zh(doc, "    - R(t)：t时刻后的记忆保持率")
    add_paragraph_zh(doc, "    - t：距上次教学交互的时间（天）")
    add_paragraph_zh(doc, "    - S：记忆强度，S = S₀ + α×N^β（N为交互次数，α=0.3，β=0.5）")
    add_paragraph_zh(doc, "  个性化遗忘率：P(F) = BaseForgetRate / (1 + 0.15 × N)")
    add_paragraph_zh(doc, "    - 交互次数N越多，遗忘率越低（记忆越稳固）")
    add_paragraph_zh(doc, "  复习触发：扫描所有知识单元，当R(t) < 0.5时触发复习提醒")
    add_paragraph_zh(doc, "  SM-2+BKT融合调度：")
    add_paragraph_zh(doc, "    - 基础间隔由SM-2算法根据历史表现计算")
    add_paragraph_zh(doc, "    - 调整规则：掌握度<0.7时缩短至70%，掌握度>0.9且评级容易时延长至120%")
    add_paragraph_zh(doc, "【区别本质】静态知识假设 vs 动态遗忘-复习循环；独立调度 vs 融合调度")
    add_paragraph_zh(doc, "【创造性判断】★★★☆☆ 遗忘曲线本身公知，但与BKT掌握度融合调度具有创新性；需强调融合公式而非仅陈述遗忘曲线")
    
    # D6
    add_heading_zh(doc, "D6: 代码执行验证机制（核心区别）", 2)
    add_paragraph_zh(doc, "【对方】未涉及代码执行，停留在文本交互层面")
    add_paragraph_zh(doc, "【我方】沙箱执行+多维度评估+知识状态联动：")
    add_paragraph_zh(doc, "  隔离执行环境：")
    add_paragraph_zh(doc, "    - Docker容器化沙箱")
    add_paragraph_zh(doc, "    - 资源限制：5秒超时、128MB内存")
    add_paragraph_zh(doc, "    - 临时文件自动清理")
    add_paragraph_zh(doc, "  多维度评估：")
    add_paragraph_zh(doc, "    - 语法维度：编译/解析是否通过")
    add_paragraph_zh(doc, "    - 逻辑维度：输出结果是否符合预期")
    add_paragraph_zh(doc, "    - 效率维度：时间/空间复杂度评估")
    add_paragraph_zh(doc, "  执行结果-知识状态联动：")
    add_paragraph_zh(doc, "    - 执行成功：涉及知识单元掌握概率+10%，错误概念严重度衰减")
    add_paragraph_zh(doc, "    - 语法错误：匹配错误概念目录后标记对应错误概念")
    add_paragraph_zh(doc, "    - 逻辑错误：根据错误类型降低对应知识单元掌握概率")
    add_paragraph_zh(doc, "【区别本质】纯文本交互 vs 代码执行闭环；无反馈 vs 执行结果反向传播")
    add_paragraph_zh(doc, "【创造性判断】★★★★☆ 编程教育场景特有的执行-反馈闭环具有领域创新性")
    
    # D7
    add_heading_zh(doc, "D7: AST级代码守卫（核心区别）", 2)
    add_paragraph_zh(doc, "【对方】未涉及代码结构验证")
    add_paragraph_zh(doc, "【我方】抽象语法树级语法结构-知识单元映射：")
    add_paragraph_zh(doc, "  语法结构解析：使用AST解析器将代码解析为语法树")
    add_paragraph_zh(doc, "  结构-知识映射表（示例）：")
    add_paragraph_zh(doc, "    ast.For → [for_loop_range, for_loop_iterable, loop_variable]")
    add_paragraph_zh(doc, "    ast.While → [while_loop, loop_condition, boolean_expression]")
    add_paragraph_zh(doc, "    ast.If → [if_statement, if_else, boolean_expression, comparison_operator]")
    add_paragraph_zh(doc, "    ast.FunctionDef → [function_definition, parameters, return_statement]")
    add_paragraph_zh(doc, "  守卫验证流程：")
    add_paragraph_zh(doc, "    1) 遍历AST所有节点")
    add_paragraph_zh(doc, "    2) 通过映射表收集所需知识单元集合RequiredKU")
    add_paragraph_zh(doc, "    3) 查询已掌握知识单元集合LearnedKU")
    add_paragraph_zh(doc, "    4) 验证：RequiredKU ⊆ LearnedKU？")
    add_paragraph_zh(doc, "    5) 若验证失败，阻止执行并提示缺少的前置知识")
    add_paragraph_zh(doc, "【技术优势】")
    add_paragraph_zh(doc, "  - 精确识别语法结构类型（vs 正则表达式仅匹配关键字）")
    add_paragraph_zh(doc, "  - 识别嵌套结构复杂度（函数内的循环内的条件判断）")
    add_paragraph_zh(doc, "  - 支持多语言扩展（Python/JS/Java等不同AST解析器）")
    add_paragraph_zh(doc, "【区别本质】无代码结构验证 vs AST级精确映射与验证")
    add_paragraph_zh(doc, "【创造性判断】★★★★★ AST-知识单元映射在编程教育场景下具有显著技术性，是代码复杂度与知识边界精确匹配的关键手段")
    
    # D8
    add_heading_zh(doc, "D8: 多策略自适应任务选择", 2)
    add_paragraph_zh(doc, "【对方】未公开任务选择策略")
    add_paragraph_zh(doc, "【我方】五策略自适应切换：")
    add_paragraph_zh(doc, "  策略1-覆盖策略：优先选择掌握度<0.3的知识单元对应题目")
    add_paragraph_zh(doc, "  策略2-难度匹配策略：选择掌握度在ZPD区间[0.4,0.85]的题目")
    add_paragraph_zh(doc, "  策略3-间隔重复策略：优先选择最久未练习（>7天）的知识单元")
    add_paragraph_zh(doc, "  策略4-错误概念导向策略：优先选择与活跃错误概念相关的题目")
    add_paragraph_zh(doc, "  策略5-不确定性最大化策略：选择掌握度最接近0.5的题目")
    add_paragraph_zh(doc, "    - 公式：score = 1.0 - min(|P(Know) - 0.5|)")
    add_paragraph_zh(doc, "    - 原理：掌握度接近0.5时不确定性最大，解答结果信息增益最大")
    add_paragraph_zh(doc, "  策略选择决策树：")
    add_paragraph_zh(doc, "    IF 活跃错误概念≥2 → 错误概念导向策略")
    add_paragraph_zh(doc, "    ELSE IF 整体掌握度<0.5且学习初期 → 覆盖策略")
    add_paragraph_zh(doc, "    ELSE IF 存在>7天未练习单元 → 间隔重复策略")
    add_paragraph_zh(doc, "    ELSE IF 学习者请求挑战或连续评级容易 → 难度匹配策略")
    add_paragraph_zh(doc, "    ELSE IF 需要诊断知识边界 → 不确定性最大化策略")
    add_paragraph_zh(doc, "【区别本质】无策略/随机选择 vs 五策略自适应切换")
    add_paragraph_zh(doc, "【创造性判断】★★★☆☆ 多策略选择在自适应学习中有先例，但策略切换条件与输入变量（错误概念数、掌握度分布等）的具体化具有创新性")
    
    # D9
    add_heading_zh(doc, "D9: Reflect-Respond四阶段约束响应", 2)
    add_paragraph_zh(doc, "【对方】虚拟学生基于预设特征生成响应，无知识状态约束")
    add_paragraph_zh(doc, "【我方】四阶段管道实现知识边界强制约束：")
    add_paragraph_zh(doc, "  阶段1-知识提取：")
    add_paragraph_zh(doc, "    - 输入：学生教学消息")
    add_paragraph_zh(doc, "    - 处理：LLM提取事实知识和代码片段（最多5事实+2代码）")
    add_paragraph_zh(doc, "    - 输出：结构化提取结果JSON")
    add_paragraph_zh(doc, "  阶段2-反思存储更新：")
    add_paragraph_zh(doc, "    - 维护反思存储（reflection_store）")
    add_paragraph_zh(doc, "    - 滑动窗口：保留最近30事实+15代码片段")
    add_paragraph_zh(doc, "    - 合并新知识，超出上限移除最旧条目")
    add_paragraph_zh(doc, "  阶段3-知识检索：")
    add_paragraph_zh(doc, "    - LLM判断反思存储与当前输入的相关性")
    add_paragraph_zh(doc, "    - 返回相关条目ID列表")
    add_paragraph_zh(doc, "    - 回退机制：LLM失败时使用关键词匹配")
    add_paragraph_zh(doc, "  阶段4-约束响应生成：")
    add_paragraph_zh(doc, "    - 构建严格边界提示词：")
    add_paragraph_zh(doc, "      Strict: You may ONLY use the following knowledge in your response.")
    add_paragraph_zh(doc, "      Do not add external knowledge.")
    add_paragraph_zh(doc, "      Relevant facts you may use: [fact_list]")
    add_paragraph_zh(doc, "      Code examples: [code_list]")
    add_paragraph_zh(doc, "      Learned units: [learned_ku_ids]")
    add_paragraph_zh(doc, "      Active misconceptions: [misconception_list]")
    add_paragraph_zh(doc, "    - 强制LLM仅使用已检索知识生成响应")
    add_paragraph_zh(doc, "【区别本质】预设角色驱动响应 vs 知识状态强制约束响应")
    add_paragraph_zh(doc, "【创造性判断】★★☆☆☆ 单独看易被归为提示工程；与状态数据库+AST+执行反馈强绑定时创造性增强")
    
    # D10
    add_heading_zh(doc, "D10: 完整证据链追踪", 2)
    add_paragraph_zh(doc, "【对方】未公开证据链或审计机制")
    add_paragraph_zh(doc, "【我方】十种事件类型构建完整因果链：")
    add_paragraph_zh(doc, "  事件类型：")
    add_paragraph_zh(doc, "    1. teaching_event（教学事件）")
    add_paragraph_zh(doc, "    2. knowledge_state_update（知识状态更新）")
    add_paragraph_zh(doc, "    3. learner_dialogue（学习者对话）")
    add_paragraph_zh(doc, "    4. task_selection（任务选择）")
    add_paragraph_zh(doc, "    5. ta_attempt（TA尝试）")
    add_paragraph_zh(doc, "    6. evaluation_result（评估结果）")
    add_paragraph_zh(doc, "    7. mastery_update（掌握度更新）")
    add_paragraph_zh(doc, "    8. misconception_activation（错误概念激活）")
    add_paragraph_zh(doc, "    9. correction_event（纠正事件）")
    add_paragraph_zh(doc, "    10. relearning_event（再学习事件）")
    add_paragraph_zh(doc, "  事件记录结构：")
    add_paragraph_zh(doc, "    - event_id：全局唯一标识符")
    add_paragraph_zh(doc, "    - event_type：枚举十种类型之一")
    add_paragraph_zh(doc, "    - timestamp：ISO格式精确到毫秒")
    add_paragraph_zh(doc, "    - related_events：关联事件ID列表（构建因果图）")
    add_paragraph_zh(doc, "    - pre_state_snapshot：前置知识状态JSON快照")
    add_paragraph_zh(doc, "    - post_state_snapshot：后置知识状态JSON快照")
    add_paragraph_zh(doc, "    - payload：类型特定的详细数据")
    add_paragraph_zh(doc, "【区别本质】孤立记录 vs 完整因果链可追溯")
    add_paragraph_zh(doc, "【创造性判断】★★☆☆☆ 日志系统常规；与每次状态转移因果绑定时创新性增强")
    
    # D11
    add_heading_zh(doc, "D11: 智能模式切换", 2)
    add_paragraph_zh(doc, "【对方】未涉及模式切换机制")
    add_paragraph_zh(doc, "【我方】周期性苏格拉底式提问：")
    add_paragraph_zh(doc, "  触发条件：(conversation_message_count + 1) % 3 == 0")
    add_paragraph_zh(doc, "    - 即每第3、6、9...轮自动切换至提问模式")
    add_paragraph_zh(doc, "  问题生成：")
    add_paragraph_zh(doc, "    - 使用苏格拉底式提问模板（Why/How类型）")
    add_paragraph_zh(doc, "    - 示例：Why does this work? How is this different from X?")
    add_paragraph_zh(doc, "  回退机制：LLM失败时使用通用模板Why is that? Can you give me an example?")
    add_paragraph_zh(doc, "【区别本质】单一对话模式 vs 周期性提问模式切换")
    add_paragraph_zh(doc, "【创造性判断】★☆☆☆☆ 教学策略常规实现，建议作为从属或说明书效果，不单独作为核心区别点")
    
    # D12
    add_heading_zh(doc, "D12: 教学解释多维质量评估", 2)
    add_paragraph_zh(doc, "【对方】未公开教学质量评估机制")
    add_paragraph_zh(doc, "【我方】三维度质量评分：")
    add_paragraph_zh(doc, "  评估维度：")
    add_paragraph_zh(doc, "    1. 准确性评分（0-1）：教学内容的事实正确性")
    add_paragraph_zh(doc, "    2. 完整性评分（0-1）：是否涵盖知识单元关键要素")
    add_paragraph_zh(doc, "    3. 清晰度评分（0-1）：表达清晰易懂程度（基于语言复杂度）")
    add_paragraph_zh(doc, "  质量评分影响BKT更新：")
    add_paragraph_zh(doc, "    - 标准习得：P(L_new) = P(L) + (1-P(L)) × P(T)")
    add_paragraph_zh(doc, "    - 质量加权：P(L_new) = P(L) + (1-P(L)) × P(T) × QualityScore × 1.2")
    add_paragraph_zh(doc, "    - 高质量教学（QualityScore>0.8）额外增加20%习得概率加成")
    add_paragraph_zh(doc, "【区别本质】无质量评估 vs 多维质量加权BKT更新")
    add_paragraph_zh(doc, "【创造性判断】★★☆☆☆ 质量评估本身常规；与BKT习得概率联动具有一定创新性")
    doc.add_paragraph()

    # Section 4: 创造性综合评估
    add_heading_zh(doc, "四、创造性综合评估矩阵", 1)
    
    add_paragraph_zh(doc, "【创造性最强】★★★★★（可作为核心答辩点）")
    add_paragraph_zh(doc, "  - D4 错误概念生命周期管理：形成-固化-检测-纠正-衰减五阶段完整机制")
    add_paragraph_zh(doc, "  - D7 AST级代码守卫：语法结构-知识单元精确映射（编程教育特有）")
    
    add_paragraph_zh(doc, "【创造性较强】★★★★☆（重要支撑点）")
    add_paragraph_zh(doc, "  - D2 动态能力构建：从零开始的细粒度知识单元演化模型")
    add_paragraph_zh(doc, "  - D3 六参数BKT：新增错误形成/纠正参数，公式具体化")
    add_paragraph_zh(doc, "  - D6 代码执行闭环：执行结果反向传播至知识状态")
    
    add_paragraph_zh(doc, "【创造性中等】★★★☆☆（需与其他点协同）")
    add_paragraph_zh(doc, "  - D5 遗忘-复习融合：强调融合调度公式而非仅遗忘曲线")
    add_paragraph_zh(doc, "  - D8 多策略选题：策略切换条件与输入变量具体化")
    add_paragraph_zh(doc, "  - D12 质量评估：强调与BKT习得概率联动")
    
    add_paragraph_zh(doc, "【创造性较弱】★★☆☆☆或★☆☆☆☆（不宜单独作为主创新）")
    add_paragraph_zh(doc, "  - D1 应用场景：需绑定技术实现差异")
    add_paragraph_zh(doc, "  - D9 Reflect-Respond：易被归为提示工程，需绑定状态库+AST")
    add_paragraph_zh(doc, "  - D10 证据链：日志系统常规，需强调与状态转移因果绑定")
    add_paragraph_zh(doc, "  - D11 模式切换：教学策略常规")
    doc.add_paragraph()

    # Section 5: 审查风险与应对
    add_heading_zh(doc, "五、可能的审查意见预判及应对", 1)
    
    add_heading_zh(doc, "预判1：「以教促学+LLM虚拟学生」被其他文献公开", 2)
    add_paragraph_zh(doc, "风险：可能存在其他专利或论文描述类似教学关系")
    add_paragraph_zh(doc, "应对：不将「逆向教学」作为主创新点，重点强调技术实现差异（D2-D8的技术闭环）")
    
    add_heading_zh(doc, "预判2：「BKT、遗忘曲线、沙箱」各自为公知常识", 2)
    add_paragraph_zh(doc, "风险：容易被认定为常规组合")
    add_paragraph_zh(doc, "应对：强调组合后解决的具体技术问题：")
    add_paragraph_zh(doc, "  - 防止LLM超越当前知识边界输出（D9+D2+D3协同）")
    add_paragraph_zh(doc, "  - 编程结果与知识状态一致（D7+D6协同）")
    add_paragraph_zh(doc, "  - 复习间隔随掌握度自适应（D5+D3协同）")
    
    add_heading_zh(doc, "预判3：「提示词/管道」创造性不足", 2)
    add_paragraph_zh(doc, "风险：D9单独支撑弱，易被质疑为提示工程")
    add_paragraph_zh(doc, "应对：权利要求中体现数据输入来源（状态矩阵+AST允许集合+沙箱反馈）与输出用途（非泛泛「优化对话」）")
    
    add_heading_zh(doc, "预判4：对比文件「学习状态信息」涵盖我方状态更新", 2)
    add_paragraph_zh(doc, "风险：审查员可能认为对比文件已公开广义状态更新")
    add_paragraph_zh(doc, "应对：用D3/D4/D7的数据结构、状态种类（五状态vs二元）、更新触发源（代码执行/错误模式/遗忘曲线）细化差异")
    doc.add_paragraph()

    # Section 6: 建议的答复/撰写策略
    add_heading_zh(doc, "六、建议的专利答复/撰写策略", 1)
    
    add_heading_zh(doc, "策略1：以「编程教育场景的技术闭环」为核心答辩线", 2)
    add_paragraph_zh(doc, "主线：AST语法守卫（D7）→ 沙箱执行（D6）→ 执行结果反向传播（D3/D4）→ 知识状态约束响应（D2/D9）")
    add_paragraph_zh(doc, "辅助：遗忘-复习融合（D5）、多策略选题（D8）、质量评估（D12）")
    add_paragraph_zh(doc, "弱化：应用场景（D1）、证据链（D10）、模式切换（D11）")
    
    add_heading_zh(doc, "策略2：权利要求布局建议", 2)
    add_paragraph_zh(doc, "独立权利要求1（系统）：包含D2+D3+D4+D6+D7+D9的核心模块组合")
    add_paragraph_zh(doc, "独立权利要求2（方法）：16步完整流程，包含遗忘计算、AST守卫、模式切换、证据链补全等对比文件没有的技术步骤")
    add_paragraph_zh(doc, "从属权利要求：将六参数公式、AST映射表、策略决策树、SM-2+BKT融合规则、约束提示格式等具体化")
    
    add_heading_zh(doc, "策略3：说明书撰写重点", 2)
    add_paragraph_zh(doc, "1) 详细描述与对比文件CN119624716A的技术缺陷对应关系（背景技术部分）")
    add_paragraph_zh(doc, "2) 提供9个详细实施例，每个区别点至少1个实施例支撑")
    add_paragraph_zh(doc, "3) 所有公式、映射表、决策树用具体数值示例说明")
    add_paragraph_zh(doc, "4) 强调模块间的数据流与控制关系，避免孤立描述")
    
    add_heading_zh(doc, "策略4：证据准备", 2)
    add_paragraph_zh(doc, "1) 系统架构图1张（展示8大模块交互）")
    add_paragraph_zh(doc, "2) 主流程图1张（教学→状态更新→约束生成→选题→AST→沙箱→回写）")
    add_paragraph_zh(doc, "3) 与对比文件权利要求1的feature-by-feature对照表（代理可用于意见陈述）")
    add_paragraph_zh(doc, "4) （可选）实验数据：六参数BKT vs 四参数BKT的预测准确率对比")
    doc.add_paragraph()

    # Conclusion
    add_heading_zh(doc, "七、一句话结论", 1)
    add_paragraph_zh(doc, "与CN119624716A最易拉开距离的核心区别：")
    add_paragraph_zh(doc, "  【技术闭环】AST语法守卫 + 沙箱执行驱动知识状态/错误概念 + 动态知识边界约束生成")
    add_paragraph_zh(doc, "  【关键创新】六参数BKT错误概念生命周期 + 语法结构-知识单元精确映射")
    add_paragraph_zh(doc, "单独支撑创造性较弱：应用场景差异、纯提示工程、常规日志系统")
    add_paragraph_zh(doc, "建议答复重点：D3+D4+D6+D7+D2的技术协同，而非单独的「逆向教学」叙事")
    add_paragraph_zh(doc, "是否具备创造性：须充分检索后定论；从交底角度，上述协同具备较强的创造性答辩基础")
    doc.add_paragraph()

    # Footer
    add_paragraph_zh(doc, "---", indent=False)
    add_paragraph_zh(doc, "文档版本：与增强版说明书对应整理", indent=False)
    add_paragraph_zh(doc, "注意：若实际产品实现与文字有出入，请以代码与真实流程为准并同步修订", indent=False)

    # Save with English filename
    output_path = Path(__file__).resolve().parent / "Patent_Comparison_Deep_Analysis_vs_CN119624716A.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    main()
