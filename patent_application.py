# -*- coding: utf-8 -*-
"""
中国大陆发明专利申请文档生成器
项目名称：基于逆向教学范式的智能计算机科学教育系统
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold


def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    # 使用英文字体名称避免编码问题
    p = doc.add_paragraph()
    run = p.add_run(text)
    font_size = 16 if level == 1 else (14 if level == 2 else 12)
    set_chinese_font(run, 'SimHei', font_size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph_zh(doc, text, indent=True, bold=False, font_size=12):
    """添加中文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
    run = p.add_run(text)
    set_chinese_font(run, '宋体', font_size, bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def create_patent_application():
    """创建专利申请文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # ==================== 说明书 ====================
    # 说明书标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('发明专利申请文件')
    set_chinese_font(run, '黑体', 22, bold=True)
    title.paragraph_format.space_after = Pt(24)
    
    # 发明名称
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run('发明名称：基于逆向教学范式的智能计算机科学教育系统')
    set_chinese_font(run, '黑体', 16, bold=True)
    name_para.paragraph_format.space_after = Pt(24)
    
    # 分隔线
    doc.add_paragraph('─' * 40)
    
    # 技术领域
    add_heading_zh(doc, '技术领域', 1)
    add_paragraph_zh(doc, 
        '本发明涉及人工智能教育技术领域，具体涉及一种基于逆向教学范式的智能计算机科学教育系统，'
        '尤其是一种通过"学生教授AI智能体"模式实现编程知识学习的创新型在线教育平台。')
    
    # 背景技术
    add_heading_zh(doc, '背景技术', 1)
    add_paragraph_zh(doc,
        '传统的计算机科学教育主要采用"教师-学生"单向传授模式，即教师向学生讲解编程概念、'
        '演示代码示例，学生被动接受知识。近年来，虽然在线教育平台和AI辅助教学工具得到了广泛应用，'
        '但仍存在以下技术问题：')
    
    add_paragraph_zh(doc,
        '1. 传统教学模式中学生参与度低，学习动机不足，难以实现深度学习；')
    add_paragraph_zh(doc,
        '2. 现有AI教学系统多采用AI作为教师、学生作为学习者的模式，忽视了"通过教授他人来学习"'
        '这一被教育心理学验证的高效学习策略；')
    add_paragraph_zh(doc,
        '3. 现有系统缺乏对学生知识掌握状态的精确追踪，无法提供个性化的教学反馈；')
    add_paragraph_zh(doc,
        '4. 现有教育平台难以模拟真实的学习过程，包括知识遗忘、错误概念形成与纠正等认知现象。')
    
    add_paragraph_zh(doc,
        '因此，亟需一种能够充分发挥学生主动性、基于科学学习理论、具备智能知识追踪能力的'
        '新型计算机科学教育系统。')
    
    # 发明内容
    add_heading_zh(doc, '发明内容', 1)
    add_paragraph_zh(doc,
        '本发明的目的在于提供一种基于逆向教学范式的智能计算机科学教育系统，解决现有技术中'
        '学生参与度低、缺乏个性化教学反馈、难以实现深度学习的技术问题。')
    
    add_paragraph_zh(doc,
        '为实现上述目的，本发明采用如下技术方案：')
    
    add_paragraph_zh(doc,
        '一种基于逆向教学范式的智能计算机科学教育系统，包括：')
    
    add_paragraph_zh(doc,
        '学生端交互模块，用于接收学生输入的教学内容，展示AI智能体的学习反馈；')
    
    add_paragraph_zh(doc,
        'AI智能体模块，基于大语言模型构建，初始状态模拟零编程知识的初学者，'
        '用于接收学生的教学输入并生成学习响应；')
    
    add_paragraph_zh(doc,
        '知识状态追踪模块，采用贝叶斯知识追踪算法，用于实时追踪AI智能体对各知识单元的掌握状态，'
        '所述知识状态包括未知、部分掌握、已掌握、存在错误概念、已纠正五种状态；')
    
    add_paragraph_zh(doc,
        '错误概念模拟模块，用于模拟初学者在学习编程过程中常见的错误概念，包括但不限于语法错误理解、'
        '概念混淆、逻辑误区等；')
    
    add_paragraph_zh(doc,
        'Reflect-Respond响应生成管道，用于在处理学生教学输入时，先进行内部反思，'
        '再根据反思结果生成符合当前知识状态的响应；')
    
    add_paragraph_zh(doc,
        '遗忘曲线建模模块，基于艾宾浩斯遗忘曲线理论，用于模拟AI智能体的知识遗忘过程，'
        '并根据练习次数调整个性化遗忘速率；')
    
    add_paragraph_zh(doc,
        '自适应测试模块，用于根据AI智能体的知识状态动态生成测试题目，评估学习效果；')
    
    add_paragraph_zh(doc,
        '代码沙箱模块，用于安全执行AI智能体生成的代码，验证其正确性；')
    
    add_paragraph_zh(doc,
        '多领域适配器模块，用于支持不同计算机科学领域的教学，包括Python编程、数据库、AI素养等。')
    
    add_paragraph_zh(doc, '优选的，所述知识状态追踪模块还包括：')
    add_paragraph_zh(doc,
        '前置依赖检查单元，用于在更新知识状态前验证知识单元间的依赖关系，确保学习路径符合认知规律；')
    add_paragraph_zh(doc,
        '掌握度量化单元，用于将知识掌握程度量化为概率值，范围从0到1。')
    
    add_paragraph_zh(doc, '优选的，所述Reflect-Respond响应生成管道包括：')
    add_paragraph_zh(doc,
        '反思单元，用于基于当前知识状态对教学内容进行内部分析和理解；')
    add_paragraph_zh(doc,
        '约束响应单元，用于根据知识状态限制响应内容，仅使用已掌握的知识单元进行回应。')
    
    add_paragraph_zh(doc,
        '本发明还提供一种基于逆向教学范式的计算机科学教学方法，包括以下步骤：')
    
    add_paragraph_zh(doc,
        'S1：初始化AI智能体的知识状态，将所有知识单元设为未知状态；')
    add_paragraph_zh(doc,
        'S2：接收学生输入的教学内容，记录教学事件；')
    add_paragraph_zh(doc,
        'S3：基于贝叶斯知识追踪算法更新AI智能体对各知识单元的掌握概率；')
    add_paragraph_zh(doc,
        'S4：检测教学内容中是否包含错误概念，如有则更新错误概念状态；')
    add_paragraph_zh(doc,
        'S5：通过Reflect-Respond管道生成AI智能体的学习响应，所述响应受当前知识状态约束；')
    add_paragraph_zh(doc,
        'S6：根据知识状态选择测试题目，由AI智能体尝试解决；')
    add_paragraph_zh(doc,
        'S7：在代码沙箱中执行AI智能体生成的代码，评估其正确性；')
    add_paragraph_zh(doc,
        'S8：基于遗忘曲线模型更新知识掌握度，模拟知识遗忘过程；')
    add_paragraph_zh(doc,
        'S9：生成知识掌握度报告，展示AI智能体的学习进展。')
    
    add_paragraph_zh(doc,
        '本发明的有益效果：')
    
    add_paragraph_zh(doc,
        '1. 通过逆向教学范式，将学生从被动学习者转变为主动教授者，显著提升学习动机和参与度；')
    add_paragraph_zh(doc,
        '2. 基于"通过教而学"的教育心理学原理，促进学生的深度学习和知识内化；')
    add_paragraph_zh(doc,
        '3. 通过贝叶斯知识追踪和遗忘曲线建模，实现对学习过程的精确模拟和个性化反馈；')
    add_paragraph_zh(doc,
        '4. 错误概念模拟功能帮助学生理解常见错误，提升教学效果；')
    add_paragraph_zh(doc,
        '5. 知识状态约束的LLM响应生成确保AI智能体的行为符合学习规律，增强教学真实感。')
    
    # 附图说明
    add_heading_zh(doc, '附图说明', 1)
    add_paragraph_zh(doc,
        '图1为本发明系统的整体架构示意图；')
    add_paragraph_zh(doc,
        '图2为知识状态追踪模块的工作原理示意图；')
    add_paragraph_zh(doc,
        '图3为Reflect-Respond响应生成管道的工作流程图；')
    add_paragraph_zh(doc,
        '图4为教学-测试循环的流程图；')
    add_paragraph_zh(doc,
        '图5为具体实施例中的用户界面示意图。')
    
    # 具体实施方式
    add_heading_zh(doc, '具体实施方式', 1)
    add_paragraph_zh(doc,
        '下面结合具体实施例对本发明作进一步详细说明。')
    
    add_heading_zh(doc, '实施例1：系统架构', 2)
    add_paragraph_zh(doc,
        '如图1所示，本发明的基于逆向教学范式的智能计算机科学教育系统包括以下核心组件：')
    
    add_paragraph_zh(doc,
        '前端采用React框架开发，包括学生端和教师端界面。学生端包含仪表板、教学页面、'
        '测试页面、知识掌握度可视化页面等；教师端提供学生管理、学习分析、教学数据可视化等功能。')
    
    add_paragraph_zh(doc,
        '后端采用FastAPI框架，提供RESTful API接口。核心API模块包括：认证模块、TA管理模块、'
        '教学流程模块、测试评估模块、知识状态模块、代码沙箱模块等。')
    
    add_paragraph_zh(doc,
        '数据层使用SQLite进行数据持久化，存储用户信息、知识状态、教学历史、测试记录等。')
    
    add_heading_zh(doc, '实施例2：知识状态追踪', 2)
    add_paragraph_zh(doc,
        '如图2所示，知识状态追踪模块采用贝叶斯知识追踪（Bayesian Knowledge Tracing, BKT）算法，'
        '对每个知识单元维护一个掌握概率值P(L)。')
    
    add_paragraph_zh(doc,
        '知识单元的状态包括：')
    add_paragraph_zh(doc,
        '- 未知（unknown）：P(L) < 0.3，表示该知识单元尚未学习；')
    add_paragraph_zh(doc,
        '- 部分掌握（partially_learned）：0.3 ≤ P(L) < 0.7，表示正在学习过程中；')
    add_paragraph_zh(doc,
        '- 已掌握（learned）：P(L) ≥ 0.7，表示已学会该知识单元；')
    add_paragraph_zh(doc,
        '- 存在错误概念（misconception）：检测到对该知识单元的错误理解；')
    add_paragraph_zh(doc,
        '- 已纠正（corrected）：错误概念已被纠正。')
    
    add_paragraph_zh(doc,
        '贝叶斯更新公式如下：')
    add_paragraph_zh(doc,
        'P(L_n|Correct) = P(L_n) + (1 - P(L_n)) * P(T)')
    add_paragraph_zh(doc,
        '其中P(L_n)为先验掌握概率，P(T)为学习概率。')
    
    add_paragraph_zh(doc,
        '遗忘曲线建模基于艾宾浩斯遗忘曲线，遗忘速率随练习次数增加而降低：')
    add_paragraph_zh(doc,
        'R(t) = e^(-t/S)，其中S为记忆强度，与练习次数正相关。')
    
    add_heading_zh(doc, '实施例3：Reflect-Respond管道', 2)
    add_paragraph_zh(doc,
        '如图3所示，AI智能体的响应生成采用Reflect-Respond两阶段管道：')
    
    add_paragraph_zh(doc,
        '反思阶段：大语言模型首先分析学生教学内容，结合当前知识状态进行内部理解。'
        '系统提示词中明确告知模型当前已掌握和未掌握的知识单元。')
    
    add_paragraph_zh(doc,
        '响应阶段：基于反思结果生成响应，响应内容严格受知识状态约束——只使用已掌握的知识单元，'
        '对未掌握内容表现出困惑或请求进一步解释，对存在错误概念的内容表现出相应的错误理解。')
    
    add_heading_zh(doc, '实施例4：教学-测试循环', 2)
    add_paragraph_zh(doc,
        '如图4所示，完整教学流程如下：')
    
    add_paragraph_zh(doc,
        '学生通过聊天界面输入教学内容（如解释Python的for循环概念）。系统记录教学事件，'
        '更新对应知识单元的掌握概率。如果教学内容包含常见错误概念，系统会标记对应的错误概念状态。')
    
    add_paragraph_zh(doc,
        'AI智能体通过Reflect-Respond管道生成响应，模拟初学者的学习过程。响应可能包括：'
        '理解反馈、提问请求澄清、展示对概念的部分理解等。')
    
    add_paragraph_zh(doc,
        '系统根据当前知识状态从题库中选择适当难度的编程题目。AI智能体尝试生成解决方案代码，'
        '代码在安全的沙箱环境中执行，评估其正确性。')
    
    add_paragraph_zh(doc,
        '根据测试结果，进一步更新知识状态。正确回答提升掌握概率，错误回答标记可能的错误概念。')
    
    add_paragraph_zh(doc,
        '定期基于遗忘曲线模型更新知识掌握度，模拟真实学习中的遗忘现象，鼓励学生进行复习教学。')
    
    add_heading_zh(doc, '实施例5：多领域支持', 2)
    add_paragraph_zh(doc,
        '系统采用领域适配器模式支持多个计算机科学领域：')
    
    add_paragraph_zh(doc,
        'Python编程领域：包含基础语法、控制流、函数、数据结构等知识单元，以及对应的练习题库。')
    add_paragraph_zh(doc,
        '数据库领域：包含SQL语法、数据库设计、查询优化等知识单元。')
    add_paragraph_zh(doc,
        'AI素养领域：包含机器学习基础、神经网络概念、AI伦理等知识单元。')
    
    add_paragraph_zh(doc,
        '每个领域独立定义知识单元、前置依赖关系、错误概念目录、评估标准和提示词模板。')
    
    # 分页
    doc.add_page_break()
    
    # ==================== 权利要求书 ====================
    rights_title = doc.add_paragraph()
    rights_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rights_title.add_run('权利要求书')
    set_chinese_font(run, '黑体', 18, bold=True)
    rights_title.paragraph_format.space_after = Pt(18)
    
    # 权利要求1 - 独立权利要求
    p = doc.add_paragraph()
    run = p.add_run('1. 一种基于逆向教学范式的智能计算机科学教育系统，其特征在于，包括：')
    set_chinese_font(run, '宋体', 12)
    
    add_paragraph_zh(doc, 
        '学生端交互模块，用于接收学生输入的教学内容，展示AI智能体的学习反馈；', 
        indent=True)
    add_paragraph_zh(doc, 
        'AI智能体模块，基于大语言模型构建，初始状态模拟零编程知识的初学者，用于接收学生的教学输入并生成学习响应；', 
        indent=True)
    add_paragraph_zh(doc, 
        '知识状态追踪模块，采用贝叶斯知识追踪算法，用于实时追踪AI智能体对各知识单元的掌握状态，所述知识状态包括未知、部分掌握、已掌握、存在错误概念、已纠正五种状态；', 
        indent=True)
    add_paragraph_zh(doc, 
        '错误概念模拟模块，用于模拟初学者在学习编程过程中常见的错误概念；', 
        indent=True)
    add_paragraph_zh(doc, 
        'Reflect-Respond响应生成管道，用于在处理学生教学输入时，先进行内部反思，再根据反思结果生成符合当前知识状态的响应；', 
        indent=True)
    add_paragraph_zh(doc, 
        '遗忘曲线建模模块，基于艾宾浩斯遗忘曲线理论，用于模拟AI智能体的知识遗忘过程，并根据练习次数调整个性化遗忘速率。', 
        indent=True)
    
    # 权利要求2
    add_paragraph_zh(doc, 
        '2. 根据权利要求1所述的系统，其特征在于，所述知识状态追踪模块还包括：\n'
        '前置依赖检查单元，用于在更新知识状态前验证知识单元间的依赖关系，确保学习路径符合认知规律；\n'
        '掌握度量化单元，用于将知识掌握程度量化为概率值，范围从0到1。',
        indent=False)
    
    # 权利要求3
    add_paragraph_zh(doc, 
        '3. 根据权利要求1所述的系统，其特征在于，所述Reflect-Respond响应生成管道包括：\n'
        '反思单元，用于基于当前知识状态对教学内容进行内部分析和理解；\n'
        '约束响应单元，用于根据知识状态限制响应内容，仅使用已掌握的知识单元进行回应。',
        indent=False)
    
    # 权利要求4
    add_paragraph_zh(doc, 
        '4. 根据权利要求1所述的系统，其特征在于，还包括自适应测试模块，用于根据AI智能体的知识状态动态生成测试题目，评估学习效果。',
        indent=False)
    
    # 权利要求5
    add_paragraph_zh(doc, 
        '5. 根据权利要求1所述的系统，其特征在于，还包括代码沙箱模块，用于安全执行AI智能体生成的代码，验证其正确性。',
        indent=False)
    
    # 权利要求6
    add_paragraph_zh(doc, 
        '6. 根据权利要求1所述的系统，其特征在于，还包括多领域适配器模块，用于支持不同计算机科学领域的教学，包括Python编程、数据库、AI素养等。',
        indent=False)
    
    # 权利要求7
    add_paragraph_zh(doc, 
        '7. 根据权利要求6所述的系统，其特征在于，每个领域独立定义知识单元、前置依赖关系、错误概念目录、评估标准和提示词模板。',
        indent=False)
    
    # 权利要求8 - 方法权利要求
    add_paragraph_zh(doc, 
        '8. 一种基于逆向教学范式的计算机科学教学方法，其特征在于，包括以下步骤：',
        indent=False)
    add_paragraph_zh(doc, 'S1：初始化AI智能体的知识状态，将所有知识单元设为未知状态；', indent=True)
    add_paragraph_zh(doc, 'S2：接收学生输入的教学内容，记录教学事件；', indent=True)
    add_paragraph_zh(doc, 'S3：基于贝叶斯知识追踪算法更新AI智能体对各知识单元的掌握概率；', indent=True)
    add_paragraph_zh(doc, 'S4：检测教学内容中是否包含错误概念，如有则更新错误概念状态；', indent=True)
    add_paragraph_zh(doc, 'S5：通过Reflect-Respond管道生成AI智能体的学习响应，所述响应受当前知识状态约束；', indent=True)
    add_paragraph_zh(doc, 'S6：根据知识状态选择测试题目，由AI智能体尝试解决；', indent=True)
    add_paragraph_zh(doc, 'S7：在代码沙箱中执行AI智能体生成的代码，评估其正确性；', indent=True)
    add_paragraph_zh(doc, 'S8：基于遗忘曲线模型更新知识掌握度，模拟知识遗忘过程；', indent=True)
    add_paragraph_zh(doc, 'S9：生成知识掌握度报告，展示AI智能体的学习进展。', indent=True)
    
    # 权利要求9
    add_paragraph_zh(doc, 
        '9. 根据权利要求8所述的方法，其特征在于，步骤S3中贝叶斯知识追踪算法的更新公式为：\n'
        'P(L_n|Correct) = P(L_n) + (1 - P(L_n)) * P(T)，\n'
        '其中P(L_n)为先验掌握概率，P(T)为学习概率。',
        indent=False)
    
    # 权利要求10
    add_paragraph_zh(doc, 
        '10. 根据权利要求8所述的方法，其特征在于，步骤S8中遗忘曲线模型基于艾宾浩斯遗忘曲线，'
        '遗忘速率随练习次数增加而降低。',
        indent=False)
    
    # 分页
    doc.add_page_break()
    
    # ==================== 摘要 ====================
    abstract_title = doc.add_paragraph()
    abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = abstract_title.add_run('摘  要')
    set_chinese_font(run, '黑体', 18, bold=True)
    abstract_title.paragraph_format.space_after = Pt(18)
    
    abstract_text = (
        '本发明公开了一种基于逆向教学范式的智能计算机科学教育系统及方法。系统包括学生端交互模块、'
        'AI智能体模块、知识状态追踪模块、错误概念模拟模块、Reflect-Respond响应生成管道和遗忘曲线建模模块。'
        'AI智能体基于大语言模型构建，初始状态模拟零编程知识的初学者；知识状态追踪模块采用贝叶斯知识追踪算法，'
        '实时追踪AI智能体对各知识单元的掌握状态；Reflect-Respond管道在处理教学输入时先进行内部反思，'
        '再根据当前知识状态生成约束响应。本发明通过逆向教学范式将学生从被动学习者转变为主动教授者，'
        '基于"通过教而学"的教育心理学原理促进深度学习，实现对学习过程的精确模拟和个性化反馈。'
    )
    add_paragraph_zh(doc, abstract_text, indent=True)
    
    # 关键词
    keywords = doc.add_paragraph()
    run = keywords.add_run('关键词：')
    set_chinese_font(run, '黑体', 12, bold=True)
    run = keywords.add_run('逆向教学；AI智能体；贝叶斯知识追踪；遗忘曲线；计算机科学教育')
    set_chinese_font(run, '宋体', 12)
    keywords.paragraph_format.first_line_indent = Cm(0.74)
    
    # 保存文档
    output_path = 'e:\\cs teachable agent\\发明专利申请文件-基于逆向教学范式的智能计算机科学教育系统.docx'
    doc.save(output_path)
    print(f'专利申请文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    create_patent_application()
