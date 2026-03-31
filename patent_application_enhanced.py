# -*- coding: utf-8 -*-
"""
中国大陆发明专利申请文档生成器 - 增强版
项目名称：基于六参数知识追踪与语法结构约束的编程教育智能体系统

增强重点（基于深入代码分析）：
1. 六参数贝叶斯知识追踪（6-Param BKT）- 增加错误概念形成/纠正概率
2. AST级代码守卫（抽象语法树级别的知识单元映射）
3. 多策略自适应任务选择算法（5种策略动态切换）
4. 完整证据链追踪系统（10种事件类型全程记录）
5. SM-2与BKT融合算法（间隔重复与知识状态深度融合）
6. Reflect-Respond四阶段管道（详细技术实现）
7. 模式切换机制（对话模式与提问者模式智能切换）
8. 教学解释多维质量评估
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold


def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    font_size = 16 if level == 1 else (14 if level == 2 else 12)
    set_chinese_font(run, 'SimHei', font_size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph_zh(doc, text, indent=True, bold=False, font_size=12):
    """添加中文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', font_size, bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def create_patent_application():
    """创建专利申请文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    
    # ==================== 说明书 ====================
    # 说明书标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('发明专利申请文件')
    set_chinese_font(run, 'SimHei', 22, bold=True)
    title.paragraph_format.space_after = Pt(24)
    
    # 发明名称
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run('发明名称：基于六参数知识追踪与语法结构约束的编程教育智能体系统')
    set_chinese_font(run, 'SimHei', 16, bold=True)
    name_para.paragraph_format.space_after = Pt(24)
    
    # 分隔线
    doc.add_paragraph('─' * 40)
    
    # ==================== 技术领域 ====================
    add_heading_zh(doc, '技术领域', 1)
    add_paragraph_zh(doc, 
        '本发明涉及人工智能教育技术领域，尤其涉及一种基于六参数知识追踪与语法结构约束的编程教育智能体系统，'
        '通过扩展的贝叶斯知识追踪算法、抽象语法树级别的代码结构守卫、多策略自适应任务选择、'
        '完整证据链追踪等技术手段，实现对编程学习者认知过程的高保真模拟与动态约束。')
    
    # ==================== 背景技术 ====================
    add_heading_zh(doc, '背景技术', 1)
    
    add_paragraph_zh(doc,
        '近年来，基于大语言模型的智能教育技术得到了快速发展。现有技术如CN119624716A公开了一种'
        '基于大语言模型的虚拟教学训练方法，该方法通过预设虚拟学生的特征参数构建具有特定能力的虚拟学生。'
        '此外，贝叶斯知识追踪（BKT）算法作为知识状态建模的经典方法，在智能教育领域得到了广泛应用。')
    
    add_paragraph_zh(doc,
        '然而，现有技术存在以下技术缺陷：')
    
    add_paragraph_zh(doc,
        '1. 知识追踪模型过于简化：传统BKT仅使用习得概率、遗忘概率、猜测概率、失误概率四个参数，'
        '未考虑学习者在掌握知识前必然经历的错误概念形成与纠正过程，无法模拟真实学习中的认知冲突与概念重构；')
    
    add_paragraph_zh(doc,
        '2. 代码安全验证机制粗放：现有编程教育系统多采用正则表达式或关键词匹配进行代码过滤，'
        '无法精确识别代码使用的语法结构，难以将代码复杂度与学习者知识状态进行精确匹配；')
    
    add_paragraph_zh(doc,
        '3. 任务选择策略单一：现有系统多采用固定顺序或随机选择方式分配练习题目，'
        '未考虑学习者当前知识状态的多样性需求（如覆盖薄弱点、维持学习流、纠正错误概念等），'
        '无法实现自适应的个性化学习路径规划；')
    
    add_paragraph_zh(doc,
        '4. 学习过程缺乏可追溯性：现有系统未建立完整的证据链追踪机制，无法记录从教学输入到知识状态变更、'
        '再到代码执行反馈的全流程因果关系，难以实现学习过程的可解释性与故障诊断；')
    
    add_paragraph_zh(doc,
        '5. 间隔重复与知识状态脱节：现有间隔重复系统（如SM-2算法）独立于知识状态追踪系统运行，'
        '未将贝叶斯掌握概率与复习调度策略深度融合，无法根据知识掌握度的动态变化调整复习间隔；')
    
    add_paragraph_zh(doc,
        '6. 教学交互模式僵化：现有系统多采用单一对话模式，未根据学习进程智能切换交互策略，'
        '无法通过苏格拉底式提问激发深度学习。')
    
    add_paragraph_zh(doc,
        '因此，亟需一种能够精确追踪知识状态演进、语法结构级别约束代码生成、多策略自适应调度、'
        '全流程证据追踪的编程教育智能体系统。')
    
    # ==================== 发明内容 ====================
    add_heading_zh(doc, '发明内容', 1)
    
    add_paragraph_zh(doc,
        '本发明的目的在于提供一种基于六参数知识追踪与语法结构约束的编程教育智能体系统，'
        '解决现有技术中知识追踪模型简化、代码验证机制粗放、任务选择策略单一、学习过程不可追溯、'
        '间隔重复与知识状态脱节、教学交互模式僵化的技术问题。')
    
    add_paragraph_zh(doc,
        '为实现上述目的，本发明采用如下技术方案：')
    
    add_paragraph_zh(doc,
        '一种基于六参数知识追踪与语法结构约束的编程教育智能体系统，其特征在于，包括：')
    
    add_paragraph_zh(doc,
        '六参数贝叶斯知识追踪模块，在传统四参数BKT基础上扩展错误概念形成概率和错误概念纠正概率，'
        '形成六参数知识状态模型；所述六参数包括：习得概率P(T)、遗忘概率P(F)、猜测概率P(G)、失误概率P(S)、'
        '错误概念形成概率P(M)、错误概念纠正概率P(C)；用于为每个知识单元维护从初始未知状态经习得、固化、'
        '遗忘、错误形成、纠错到再学习的完整生命周期状态；')
    
    add_paragraph_zh(doc,
        '抽象语法树级代码守卫模块，用于建立代码语法结构到知识单元的双向映射关系，'
        '包括：语法结构解析单元，使用AST解析器将代码解析为抽象语法树；结构-知识映射单元，'
        '维护语法节点类型到知识单元标识的映射表；守卫验证单元，遍历AST节点，'
        '检查每个语法结构对应的知识单元是否处于已掌握状态，拒绝执行包含未掌握语法结构的代码；')
    
    add_paragraph_zh(doc,
        '多策略自适应任务选择模块，用于根据学习目标和知识状态动态选择题目分配策略，'
        '包括五种互斥策略：覆盖策略，优先选择掌握度低于阈值的知识单元对应题目；'
        '难度匹配策略，选择掌握度处于最近发展区[0.4,0.85]的题目；'
        '间隔重复策略，优先选择最久未练习的知识单元；'
        '错误概念导向策略，优先选择与活跃错误概念相关的题目；'
        '不确定性最大化策略，选择掌握度最接近0.5的题目以最大化信息增益；'
        '策略选择基于学习阶段、错误概念活跃度、整体掌握度分布动态决策；')
    
    add_paragraph_zh(doc,
        '完整证据链追踪模块，用于记录从教学输入到知识状态变更的全流程因果链，'
        '包括十种事件类型：教学事件、知识状态更新、学习者对话、任务选择、TA尝试、评估结果、'
        '掌握度更新、错误概念激活、纠正事件、再学习事件；每个事件记录时间戳、关联事件ID、'
        '前置知识状态、后置知识状态，形成可追溯的学习审计日志；')
    
    add_paragraph_zh(doc,
        'SM-2与BKT融合调度模块，用于将间隔重复算法与贝叶斯知识追踪深度融合，'
        '包括：基础调度单元，使用SM-2算法根据历史复习表现计算下次复习间隔；'
        '掌握度调整单元，根据BKT掌握概率动态调整间隔天数，掌握度低于0.7时缩短间隔至70%，'
        '掌握度高于0.9且评级为容易时延长间隔至120%；遗忘速率自适应单元，'
        '基于复习成功率动态调整个体遗忘速率估计值；')
    
    add_paragraph_zh(doc,
        'Reflect-Respond四阶段约束响应模块，用于根据知识状态动态约束大语言模型响应生成，'
        '包括：知识提取阶段，从教学输入中提取事实知识和代码片段；反思存储更新阶段，'
        '合并新知识至反思存储并维护滑动窗口；知识检索阶段，使用LLM判断相关性或回退至关键词匹配；'
        '约束响应生成阶段，构建严格边界提示词强制模型仅使用已检索知识生成响应；')
    
    add_paragraph_zh(doc,
        '智能模式切换模块，用于在教学对话模式与苏格拉底提问模式间智能切换，'
        '包括：计数触发单元，每累计3轮对话自动切换至提问模式一次；问题生成单元，'
        '使用启发式模板或LLM生成Why/How类型的苏格拉底式追问；回退机制，当问题生成失败时'
        '使用通用追问模板；')
    
    add_paragraph_zh(doc,
        '教学解释多维质量评估模块，用于量化评估用户教学输入的质量，'
        '包括五个评估维度：准确性评分、完整性评分、清晰度评分、错误概念纠正标记、'
        '解释类别分类；质量评分作为六参数BKT中习得概率计算的加权因子，高质量教学应用更高习得概率。')
    
    add_paragraph_zh(doc,
        '优选的，所述六参数贝叶斯知识追踪模块采用以下状态转移公式：')
    add_paragraph_zh(doc,
        '习得转移：P(L_{n+1}=1) = P(L_n) + (1-P(L_n))×P(T)×QualityScore×PrereqSatisfied')
    add_paragraph_zh(doc,
        '遗忘转移：P(L_{n+1}) = P(L_n)×(1-P(F))，其中P(F) = BaseForgetRate/(1+0.15×N)，N为交互次数')
    add_paragraph_zh(doc,
        '错误形成：若教学内容包含误导信息，以概率P(M)标记错误概念状态，P(M) = ContentMisleading×LearnerErrorProneIndex')
    add_paragraph_zh(doc,
        '错误纠正：若检测到纠错教学，以概率P(C)迁移至纠错状态，P(C) = CorrectionQuality×CognitiveConflictIntensity')
    
    add_paragraph_zh(doc,
        '优选的，所述抽象语法树级代码守卫模块的语法结构-知识单元映射表包括：')
    add_paragraph_zh(doc,
        'For节点映射至["for_loop_range", "for_loop_iterable"]知识单元；')
    add_paragraph_zh(doc,
        'While节点映射至["while_loop", "loop_condition"]知识单元；')
    add_paragraph_zh(doc,
        'If节点映射至["if_statement", "if_else", "if_elif_else", "boolean_expression"]知识单元；')
    add_paragraph_zh(doc,
        'FunctionDef节点映射至["function_definition", "parameters", "return_statement"]知识单元；'
        '守卫验证时遍历AST所有节点，收集所需知识单元集合，检查是否为已掌握知识单元的子集。')
    
    add_paragraph_zh(doc,
        '优选的，所述多策略自适应任务选择模块的策略选择逻辑为：')
    add_paragraph_zh(doc,
        '若存在活跃错误概念且数量≥2，优先采用错误概念导向策略；')
    add_paragraph_zh(doc,
        '若整体掌握度<0.5且处于学习初期，采用覆盖策略；')
    add_paragraph_zh(doc,
        '若掌握度分布均匀且存在多日未练习单元，采用间隔重复策略；')
    add_paragraph_zh(doc,
        '若学习者请求挑战或评级多为容易，采用难度匹配策略；')
    add_paragraph_zh(doc,
        '若系统需最大化信息增益诊断知识边界，采用不确定性最大化策略。')
    
    add_paragraph_zh(doc,
        '优选的，所述完整证据链追踪模块的事件记录结构包括：')
    add_paragraph_zh(doc,
        '事件ID：全局唯一标识符；')
    add_paragraph_zh(doc,
        '事件类型：枚举十种类型之一；')
    add_paragraph_zh(doc,
        '时间戳：ISO格式精确到毫秒；')
    add_paragraph_zh(doc,
        '关联事件ID列表：构建因果链图；')
    add_paragraph_zh(doc,
        '前置知识状态快照：JSON格式的完整状态矩阵；')
    add_paragraph_zh(doc,
        '后置知识状态快照：事件完成后的状态矩阵；')
    add_paragraph_zh(doc,
        '事件负载：类型特定的详细数据。')
    
    add_paragraph_zh(doc,
        '本发明还提供一种基于六参数知识追踪与语法结构约束的编程教育智能体方法，包括以下步骤：')
    
    add_paragraph_zh(doc,
        'S1：初始化阶段，构建六参数BKT状态矩阵，所有知识单元状态设为未知，'
        '六参数初始化为默认值：P(T)=0.1, P(F)=0.1, P(G)=0.25, P(S)=0.1, P(M)=0.3, P(C)=0.5；')
    
    add_paragraph_zh(doc,
        'S2：证据链记录阶段，创建教学事件记录，关联前置知识状态快照；')
    
    add_paragraph_zh(doc,
        'S3：教学质量评估阶段，从准确性、完整性、清晰度三个维度评估教学输入，'
        '检测是否包含误导信息，计算综合质量评分QualityScore；')
    
    add_paragraph_zh(doc,
        'S4：六参数BKT更新阶段，应用六参数公式更新知识单元掌握概率，'
        '检查前置依赖满足度，若不满足则抑制习得转移，若检测到误导信息则触发错误概念形成转移；')
    
    add_paragraph_zh(doc,
        'S5：遗忘计算阶段，应用个性化遗忘公式P(F)=BaseForgetRate/(1+0.15×N)，'
        '其中N为该知识单元的教学证据数，更新掌握概率，若低于阈值则标记遗忘状态；')
    
    add_paragraph_zh(doc,
        'S6：模式切换判断阶段，检查对话轮数是否满足切换条件（每3轮），'
        '若满足则生成苏格拉底式问题并输出，否则进入S7；')
    
    add_paragraph_zh(doc,
        'S7：Reflect-Respond四阶段处理阶段，依次执行知识提取、反思存储更新、知识检索、约束响应生成，'
        '输出受知识边界约束的学习者响应；')
    
    add_paragraph_zh(doc,
        'S8：任务选择阶段，根据当前错误概念活跃度、整体掌握度、复习间隔分布，'
        '从五种策略中选择适用策略，选择对应编程题目；')
    
    add_paragraph_zh(doc,
        'S9：代码生成阶段，触发AI智能体基于当前知识状态生成解决方案代码；')
    
    add_paragraph_zh(doc,
        'S10：AST守卫验证阶段，解析代码为抽象语法树，遍历所有节点，'
        '通过语法结构-知识单元映射表收集所需知识单元，检查是否为已掌握集合的子集，'
        '若验证失败则阻止执行并记录错误类型；')
    
    add_paragraph_zh(doc,
        'S11：代码执行阶段，在隔离沙箱环境中执行验证通过的代码，'
        '收集语法错误、运行时错误、逻辑错误、输出结果；')
    
    add_paragraph_zh(doc,
        'S12：执行结果-知识状态联动阶段，将执行结果反向传播至六参数BKT：'
        '执行成功则提升相关知识点掌握概率并衰减错误概念严重程度，'
        '执行失败则根据错误类型标记错误概念或降低掌握概率；')
    
    add_paragraph_zh(doc,
        'S13：证据链补全阶段，创建评估结果事件和掌握度更新事件，关联S2创建的教学事件，'
        '记录后置知识状态快照，形成完整因果链；')
    
    add_paragraph_zh(doc,
        'S14：SM-2与BKT融合调度阶段，基于当前掌握度调整SM-2计算的复习间隔，'
        '掌握度低于0.7时缩短间隔，高于0.9时延长间隔；')
    
    add_paragraph_zh(doc,
        'S15：复习触发阶段，扫描调度队列，对到期复习任务生成提醒，'
        '当用户进行复习教学时应用更高习得概率参数；')
    
    add_paragraph_zh(doc,
        'S16：循环执行S2-S15，直至达到预设学习目标。')
    
    add_paragraph_zh(doc,
        '本发明与现有技术相比的有益效果：')
    
    add_paragraph_zh(doc,
        '1. 相比传统四参数BKT，本发明的六参数模型完整覆盖知识生命周期，'
        '特别增加错误概念形成与纠正参数，实现真实学习中"正确-错误-再正确"的认知螺旋模拟；'
        '这是现有技术未涉及的重要认知过程；')
    
    add_paragraph_zh(doc,
        '2. 相比正则表达式级别的代码过滤，本发明的AST级守卫实现语法结构到知识单元的精确映射，'
        '能够识别代码使用的具体语言特性并匹配学习者知识边界，实现真正的知识状态约束代码生成；')
    
    add_paragraph_zh(doc,
        '3. 相比单一任务选择策略，本发明的五策略自适应系统根据学习阶段、错误概念状态、掌握度分布'
        '动态选择最优策略，实现覆盖薄弱点、维持学习流、纠正错误概念、最大化信息增益等多目标优化；')
    
    add_paragraph_zh(doc,
        '4. 相比孤立的学习记录，本发明的十种事件类型证据链构建完整因果图，'
        '支持从任意知识状态变更追溯至源头教学输入，实现学习过程的全透明可解释；')
    
    add_paragraph_zh(doc,
        '5. 相比独立运行的SM-2和BKT，本发明的融合调度算法根据实时掌握度动态调整复习间隔，'
        '掌握度低时增加复习频率，掌握度高时延长间隔，实现效率与效果的平衡；')
    
    add_paragraph_zh(doc,
        '6. 相比单一对话模式，本发明的模式切换机制通过周期性苏格拉底式提问激发元认知，'
        '促进深度学习，这是教育心理学验证的高效教学策略在技术实现上的创新应用。')
    
    # ==================== 附图说明 ====================
    add_heading_zh(doc, '附图说明', 1)
    add_paragraph_zh(doc, '图1为本发明系统的整体架构示意图，展示六参数BKT、AST守卫、多策略选择、证据链追踪等模块的交互关系；')
    add_paragraph_zh(doc, '图2为六参数贝叶斯知识追踪的状态转移图，展示六种概率参数驱动下的状态生命周期；')
    add_paragraph_zh(doc, '图3为抽象语法树级代码守卫的工作流程图，展示代码解析、结构映射、守卫验证的技术流程；')
    add_paragraph_zh(doc, '图4为多策略自适应任务选择的决策逻辑图，展示五种策略的选择条件与适用场景；')
    add_paragraph_zh(doc, '图5为完整证据链追踪系统的数据模型图，展示十种事件类型及其关联关系；')
    add_paragraph_zh(doc, '图6为SM-2与BKT融合调度算法的流程图，展示掌握度调整因子对复习间隔的动态修正；')
    add_paragraph_zh(doc, '图7为Reflect-Respond四阶段管道的详细流程图，展示知识提取到约束响应生成的完整过程；')
    add_paragraph_zh(doc, '图8为智能模式切换机制的状态机图，展示对话模式与提问模式的切换逻辑；')
    add_paragraph_zh(doc, '图9为具体实施例中Python编程领域的语法结构-知识单元映射表示意图。')
    
    # ==================== 具体实施方式 ====================
    add_heading_zh(doc, '具体实施方式', 1)
    add_paragraph_zh(doc, '下面结合具体实施例对本发明作进一步详细说明。')
    
    add_heading_zh(doc, '实施例1：六参数BKT的详细实现', 2)
    add_paragraph_zh(doc,
        '本实施例详细描述六参数贝叶斯知识追踪的实现方式。假设学习者正在学习Python的for循环知识单元（unit_id="for_loop"）。')
    
    add_paragraph_zh(doc,
        '初始状态：六参数初始化为默认值：P(T)=0.1（习得概率），P(F)=0.1（基础遗忘概率），'
        'P(G)=0.25（猜测概率），P(S)=0.1（失误概率），P(M)=0.3（错误形成概率），P(C)=0.5（错误纠正概率）。'
        '知识单元状态为未知，掌握概率P(L)=0.01。')
    
    add_paragraph_zh(doc,
        '教学输入处理：用户解释for循环的语法和用法。教学解释质量评估模块从三个维度评估：'
        '准确性评分0.9（解释正确），完整性评分0.8（涵盖基本用法但未涉及高级特性），清晰度评分0.85（表达清楚）。'
        '综合质量评分QualityScore = (0.9+0.8+0.85)/3 = 0.85。未检测到误导信息。')
    
    add_paragraph_zh(doc,
        '六参数更新计算：前置依赖检查确认"变量定义"和"基本运算"已掌握（满足前置条件PrereqSatisfied=1）。')
    add_paragraph_zh(doc,
        '习得转移：P(L_new) = 0.01 + (1-0.01)×0.1×0.85×1 = 0.01 + 0.084 = 0.094')
    add_paragraph_zh(doc,
        '状态迁移：掌握概率从0.01提升至0.094，虽未达习得阈值(0.7)，但已脱离纯未知状态，标记为部分习得。')
    
    add_paragraph_zh(doc,
        '错误概念形成场景：假设另一场景中用户错误地解释"for循环和while循环在任何情况下都可以互换使用"。'
        '系统检测到误导信息（ContentMisleading=1），学习者易错指数0.6。')
    add_paragraph_zh(doc,
        '错误形成概率：P(M_triggered) = 1 × 0.6 = 0.6 > 阈值0.3，触发错误概念形成。')
    add_paragraph_zh(doc,
        '系统在active_misconceptions列表中添加记录：{misconception_id: "loop_type_confusion", '
        'severity_score: 0.4, trigger_count: 1}。知识单元状态迁移至错误概念状态。')
    
    add_paragraph_zh(doc,
        '错误纠正场景：后续教学中用户纠正上述错误，详细解释两种循环的适用场景差异。'
        '纠错教学质量评估：CorrectionQuality=0.9，认知冲突强度0.8。')
    add_paragraph_zh(doc,
        '错误纠正概率：P(C) = 0.9 × 0.8 = 0.72 > 阈值0.5，触发错误纠正。')
    add_paragraph_zh(doc,
        '系统迁移状态至纠错状态，掌握概率重置为0.1（从零开始重新积累），记录纠错事件至证据链。')
    
    add_heading_zh(doc, '实施例2：AST级代码守卫的详细实现', 2)
    add_paragraph_zh(doc,
        '本实施例展示抽象语法树级代码守卫如何精确控制代码复杂度与知识状态的匹配。')
    
    add_paragraph_zh(doc,
        '语法结构-知识单元映射表定义（以Python为例）：')
    add_paragraph_zh(doc,
        "{ast.For: [for_loop_range, for_loop_iterable, loop_variable], "
        "ast.While: [while_loop, loop_condition, boolean_expression], "
        "ast.If: [if_statement, if_else, if_elif_else, boolean_expression, comparison_operator], "
        "ast.FunctionDef: [function_definition, parameters, return_statement, function_scope], "
        "ast.ListComp: [list_comprehension, for_loop_iterable, conditional_expression], "
        "ast.TryExcept: [exception_handling, try_block, except_block]}")
    
    add_paragraph_zh(doc,
        '守卫验证流程：学习者教授后，AI智能体生成以下代码：')
    add_paragraph_zh(doc,
        "def calculate_sum(numbers): total = 0; for n in numbers: if n > 0: total += n; return total")
    
    add_paragraph_zh(doc,
        'AST解析：代码解析为抽象语法树，包含以下节点类型：FunctionDef、Assign、For、If、AugAssign、Return。')
    
    add_paragraph_zh(doc,
        '所需知识单元收集：遍历AST节点，通过映射表收集所需知识单元：')
    add_paragraph_zh(doc,
        '- FunctionDef → ["function_definition", "parameters", "return_statement", "function_scope"]')
    add_paragraph_zh(doc,
        '- For → ["for_loop_range", "for_loop_iterable", "loop_variable"]')
    add_paragraph_zh(doc,
        '- If → ["if_statement", "boolean_expression", "comparison_operator"]')
    add_paragraph_zh(doc,
        '- AugAssign → ["assignment_operator", "compound_assignment"]')
    
    add_paragraph_zh(doc,
        '守卫检查：合并所有所需知识单元，去重后得到集合RequiredKU。'
        '查询六参数BKT模块，获取当前已掌握知识集合LearnedKU。'
        '验证：RequiredKU ⊆ LearnedKU？')
    
    add_paragraph_zh(doc,
        '场景A-验证通过：若学习者已掌握函数定义、for循环、if语句，则验证通过，允许执行代码。')
    
    add_paragraph_zh(doc,
        '场景B-验证失败：若学习者尚未学习"function_definition"（函数定义），则验证失败。'
        '系统阻止代码执行，记录错误类型为"使用未掌握语法结构：函数定义"，'
        '并提示"TA尚未学习如何定义函数，请先教授函数定义的基础知识"。')
    
    add_paragraph_zh(doc,
        '技术优势：相比正则表达式匹配（如仅检查"def"关键字），AST级守卫能够：')
    add_paragraph_zh(doc,
        '1. 精确识别使用的具体语法结构类型（FunctionDef vs ClassDef vs AsyncFunctionDef）；')
    add_paragraph_zh(doc,
        '2. 识别嵌套结构复杂度（函数内的循环内的条件判断）；')
    add_paragraph_zh(doc,
        '3. 支持多语言扩展（Python、JavaScript、Java的AST解析器不同，但守卫逻辑复用）。')
    
    add_heading_zh(doc, '实施例3：多策略自适应任务选择', 2)
    add_paragraph_zh(doc,
        '本实施例展示如何根据学习状态动态选择最优任务分配策略。')
    
    add_paragraph_zh(doc,
        '策略选择决策树：')
    add_paragraph_zh(doc,
        '步骤1：检查错误概念状态。若active_misconceptions列表非空且数量≥2，'
        '优先采用错误概念导向策略（Strategy_Misconception）。')
    add_paragraph_zh(doc,
        '步骤2：若无活跃错误概念，计算整体掌握度MeanP = 所有知识单元掌握概率的均值。'
        '若MeanP < 0.5且学习事件总数<50（学习初期），采用覆盖策略（Strategy_Coverage）。')
    add_paragraph_zh(doc,
        '步骤3：若MeanP ≥ 0.5，检查最久未练习时间。若存在超过7天未练习的已掌握知识单元，'
        '采用间隔重复策略（Strategy_Spaced）。')
    add_paragraph_zh(doc,
        '步骤4：若学习者最近连续3次评级为"容易"，或明确请求挑战，采用难度匹配策略（Strategy_Difficulty）。')
    add_paragraph_zh(doc,
        '步骤5：若系统需要诊断知识边界（如学习者刚完成一个知识单元的学习），'
        '采用不确定性最大化策略（Strategy_Uncertainty），选择掌握度最接近0.5的题目。')
    
    add_paragraph_zh(doc,
        '不确定性最大化策略的数学实现：')
    add_paragraph_zh(doc,
        '对于候选题目集合中的每个题目，提取其所需知识单元列表。')
    add_paragraph_zh(doc,
        '计算掌握度距0.5的距离：distance = |P(Know) - 0.5|。')
    add_paragraph_zh(doc,
        '题目得分：score = 1.0 - min(distance_list)，选择得分最高的题目。')
    add_paragraph_zh(doc,
        '原理：掌握度接近0.5意味着对该知识单元的不确定性最大，解答结果将提供最大信息增益。')
    
    add_heading_zh(doc, '实施例4：完整证据链追踪系统', 2)
    add_paragraph_zh(doc,
        '本实施例展示十种事件类型如何构建完整的学习因果链。')
    
    add_paragraph_zh(doc,
        '场景：学习者教授"Python列表推导式"概念。')
    
    add_paragraph_zh(doc,
        '事件1-教学事件（Teaching Event）：')
    add_paragraph_zh(doc,
        '  event_id: "te_001", type: "teaching_event", timestamp: "2024-03-20T10:00:00.000Z",')
    add_paragraph_zh(doc,
        '  payload: {content: "列表推导式是...", knowledge_units: ["list_comprehension"], quality_score: 0.85},')
    add_paragraph_zh(doc,
        '  pre_state_snapshot: {...}, post_state_snapshot: {...}')
    
    add_paragraph_zh(doc,
        '事件2-知识状态更新（Knowledge State Update）：')
    add_paragraph_zh(doc,
        '  event_id: "ksu_001", type: "knowledge_state_update", timestamp: "2024-03-20T10:00:01.200Z",')
    add_paragraph_zh(doc,
        '  related_events: ["te_001"], payload: {unit_id: "list_comprehension", old_p_know: 0.1, new_p_know: 0.35}')
    
    add_paragraph_zh(doc,
        '事件3-学习者对话（Learner Dialogue）：')
    add_paragraph_zh(doc,
        '  event_id: "ld_001", type: "learner_dialogue", timestamp: "2024-03-20T10:00:02.500Z",')
    add_paragraph_zh(doc,
        '  related_events: ["ksu_001"], payload: {dialogue_type: "TA_response", content: "我理解了..."}')
    
    add_paragraph_zh(doc,
        '...（后续事件链：任务选择→TA尝试→评估结果→掌握度更新→可能触发的错误概念激活/纠正/再学习）')
    
    add_paragraph_zh(doc,
        '因果追溯示例：若发现"list_comprehension"知识单元的掌握度异常降低，'
        '可通过related_events字段追溯至源头，发现是由于某次代码执行失败（RuntimeError）导致。'
        '进一步追溯该执行失败事件，发现是因为学习者使用了未掌握的"嵌套列表推导"语法。'
        '系统据此推断：需要加强"嵌套列表推导"的前置教学。')
    
    add_heading_zh(doc, '实施例5：SM-2与BKT融合调度', 2)
    add_paragraph_zh(doc,
        '本实施例展示如何将间隔重复算法与贝叶斯知识追踪深度融合。')
    
    add_paragraph_zh(doc,
        '基础SM-2计算：对于"for_loop"知识单元，上次复习评级为"良好"（Good），'
        '原间隔天数interval=3，容易度因子EF=2.5。新间隔interval = 3 × 2.5 = 7.5天。')
    
    add_paragraph_zh(doc,
        '掌握度调整层：查询六参数BKT，当前掌握度P(Know)=0.6（低于0.7阈值）。')
    add_paragraph_zh(doc,
        '调整计算：adjusted_interval = 7.5 × 0.7 = 5.25天（缩短30%）。')
    add_paragraph_zh(doc,
        '理由：虽然上次评级良好，但BKT显示掌握度尚未稳固，需增加复习频率。')
    
    add_paragraph_zh(doc,
        '场景2-延长间隔：若P(Know)=0.95（掌握度高），且评级为"容易"（Easy），')
    add_paragraph_zh(doc,
        '则adjusted_interval = 7.5 × 1.2 = 9天（延长20%）。')
    
    add_paragraph_zh(doc,
        '遗忘速率自适应：系统维护个体遗忘速率估计值forgetting_rate，初始值0.3。')
    add_paragraph_zh(doc,
        '每次成功复习：forgetting_rate = forgetting_rate × 0.9（降低，表示记忆力较好）。')
    add_paragraph_zh(doc,
        '每次失败复习：forgetting_rate = forgetting_rate × 1.1（升高，表示需要更频繁复习）。')
    
    add_heading_zh(doc, '实施例6：Reflect-Respond四阶段管道', 2)
    add_paragraph_zh(doc,
        '本实施例详细展示知识状态如何约束大语言模型的响应生成。')
    
    add_paragraph_zh(doc,
        '阶段1-知识提取：用户输入教学内容"列表推导式的语法是[expression for item in iterable]"。'
        '调用LLM API（专用提取prompt），输出JSON格式：')
    add_paragraph_zh(doc,
        '{"facts": ["列表推导式语法结构", "expression部分", "for循环部分"], '
        '"code_snippets": ["[x for x in range(10)]"]}')
    
    add_paragraph_zh(doc,
        '阶段2-反思存储更新：维护反思存储数据结构（reflection_store），包含facts和code_implementations两个列表。'
        '采用滑动窗口机制，保留最近30个事实和15个代码片段。合并新知识后，若超出上限则移除最旧条目。')
    
    add_paragraph_zh(doc,
        '阶段3-知识检索：用户新输入"请用列表推导式生成1-100的偶数列表"。'
        '调用LLM API判断反思存储中各条目与当前输入的相关性，返回相关条目ID列表。'
        '回退机制：若LLM判断失败，使用关键词匹配（如"列表推导式"关键词匹配）。')
    
    add_paragraph_zh(doc,
        '阶段4-约束响应生成：检索到相关事实：["列表推导式语法结构"]和代码片段：["[x for x in range(10)]"]。'
        '构建严格约束prompt：')
    add_paragraph_zh(doc,
        '"""Strict: You may ONLY use the following knowledge in your response. '
        'Do not add external knowledge.\n\n'
        'Relevant facts you may use:\n'
        '- 列表推导式语法结构\n\n'
        'Code examples you may reference:\n'
        '- [x for x in range(10)]\n\n'
        'Learned unit IDs: ["list_comprehension_basic"]\n'
        'Active misconceptions: []\n'
        'Unlearned units: ["list_comprehension_nested", "list_comprehension_conditional"]\n\n'
        'Based ONLY on the above knowledge, generate a response to: '
        '"请用列表推导式生成1-100的偶数列表"""')
    
    add_paragraph_zh(doc,
        '响应生成：LLM基于严格约束生成代码"[x for x in range(2, 101, 2)]"。'
        '由于约束prompt明确禁止添加外部知识，模型不会使用未掌握的"条件表达式"（if x%2==0）。'
        '而是采用已掌握的range步长参数方式。')
    
    add_heading_zh(doc, '实施例7：智能模式切换与苏格拉底提问', 2)
    add_paragraph_zh(doc,
        '本实施例展示如何通过模式切换激发元认知。')
    
    add_paragraph_zh(doc,
        '对话轮数计数：系统维护conversation_message_count计数器。')
    add_paragraph_zh(doc,
        '切换触发：当(conversation_message_count + 1) % 3 == 0时，触发模式切换。')
    add_paragraph_zh(doc,
        '即每第3、6、9...轮自动切换至提问者模式一次。')
    
    add_paragraph_zh(doc,
        '问题生成：调用LLM API，使用苏格拉底式提问模板prompt：')
    add_paragraph_zh(doc,
        '"""You are a curious learner who asks thought-provoking questions. '
        'Based on the recent teaching about {topic}, generate a Why or How question '
        'that encourages the teacher to explain the underlying principles. '
        'Examples: "Why does this work?", "How is this different from {related_concept}?", '
        '"What would happen if we changed X to Y?"""')
    
    add_paragraph_zh(doc,
        '回退机制：若LLM生成失败或超时，使用通用追问模板："Why is that? Can you give me an example?"')
    
    add_heading_zh(doc, '实施例8：教学解释多维质量评估', 2)
    add_paragraph_zh(doc,
        '本实施例展示如何量化评估教学质量并影响知识状态更新。')
    
    add_paragraph_zh(doc,
        '评估维度与评分：')
    add_paragraph_zh(doc,
        '准确性评分（0-1）：教学内容的 factual correctness。由LLM或规则引擎评估。')
    add_paragraph_zh(doc,
        '完整性评分（0-1）：是否涵盖知识单元的关键要素。通过与标准知识点列表对比。')
    add_paragraph_zh(doc,
        '清晰度评分（0-1）：表达是否清晰易懂。通过语言复杂度指标（如句子长度、术语密度）评估。')
    add_paragraph_zh(doc,
        '错误概念纠正标记（布尔）：是否识别并纠正了已存在的错误概念。')
    add_paragraph_zh(doc,
        '解释类别分类（枚举）：correct/ambiguous/misleading/partial等。')
    
    add_paragraph_zh(doc,
        '质量评分影响BKT更新：')
    add_paragraph_zh(doc,
        '标准习得转移：P(L_new) = P(L) + (1-P(L)) × P(T)')
    add_paragraph_zh(doc,
        '质量加权习得转移：P(L_new) = P(L) + (1-P(L)) × P(T) × QualityScore × 1.2')
    add_paragraph_zh(doc,
        '高质量教学（QualityScore>0.8）可额外增加20%的习得概率加成。')
    
    # 分页
    doc.add_page_break()
    
    # ==================== 权利要求书 ====================
    rights_title = doc.add_paragraph()
    rights_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rights_title.add_run('权利要求书')
    set_chinese_font(run, 'SimHei', 18, bold=True)
    rights_title.paragraph_format.space_after = Pt(18)
    
    # 权利要求1 - 独立权利要求（系统）
    p = doc.add_paragraph()
    run = p.add_run('1. 一种基于六参数知识追踪与语法结构约束的编程教育智能体系统，其特征在于，包括：')
    set_chinese_font(run, 'SimSun', 12)
    
    add_paragraph_zh(doc, 
        '六参数贝叶斯知识追踪模块，在传统四参数BKT基础上扩展错误概念形成概率和错误概念纠正概率，形成六参数知识状态模型；所述六参数包括习得概率、遗忘概率、猜测概率、失误概率、错误概念形成概率、错误概念纠正概率；用于为每个知识单元维护从初始未知状态经习得、固化、遗忘、错误形成、纠错到再学习的完整生命周期状态；', 
        indent=True)
    add_paragraph_zh(doc, 
        '抽象语法树级代码守卫模块，用于建立代码语法结构到知识单元的双向映射关系，包括语法结构解析单元、结构-知识映射单元和守卫验证单元；', 
        indent=True)
    add_paragraph_zh(doc, 
        '多策略自适应任务选择模块，用于根据学习目标和知识状态动态选择题目分配策略，包括覆盖策略、难度匹配策略、间隔重复策略、错误概念导向策略、不确定性最大化策略五种互斥策略；', 
        indent=True)
    add_paragraph_zh(doc, 
        '完整证据链追踪模块，用于记录从教学输入到知识状态变更的全流程因果链，包括十种事件类型及其关联关系；', 
        indent=True)
    add_paragraph_zh(doc, 
        'SM-2与BKT融合调度模块，用于将间隔重复算法与贝叶斯知识追踪深度融合，根据BKT掌握概率动态调整SM-2计算的复习间隔；', 
        indent=True)
    add_paragraph_zh(doc, 
        'Reflect-Respond四阶段约束响应模块，用于根据知识状态动态约束大语言模型响应生成，包括知识提取阶段、反思存储更新阶段、知识检索阶段和约束响应生成阶段。', 
        indent=True)
    
    # 权利要求2
    add_paragraph_zh(doc, 
        '2. 根据权利要求1所述的系统，其特征在于，所述六参数贝叶斯知识追踪模块的状态转移包括：习得转移公式P(L_{n+1}=1) = P(L_n) + (1-P(L_n))×P(T)×QualityScore×PrereqSatisfied；遗忘转移公式P(F) = BaseForgetRate/(1+0.15×N)，其中N为教学交互次数；错误概念形成公式P(M) = ContentMisleading×LearnerErrorProneIndex；错误概念纠正公式P(C) = CorrectionQuality×CognitiveConflictIntensity。',
        indent=False)
    
    # 权利要求3
    add_paragraph_zh(doc, 
        '3. 根据权利要求1所述的系统，其特征在于，所述抽象语法树级代码守卫模块的语法结构-知识单元映射表包括：For节点映射至for循环相关知识单元，While节点映射至while循环相关知识单元，If节点映射至条件判断相关知识单元，FunctionDef节点映射至函数定义相关知识单元；守卫验证单元遍历AST所有节点，收集所需知识单元集合，检查是否为已掌握知识单元集合的子集。',
        indent=False)
    
    # 权利要求4
    add_paragraph_zh(doc, 
        '4. 根据权利要求1所述的系统，其特征在于，所述多策略自适应任务选择模块的策略选择逻辑为：若存在活跃错误概念且数量大于等于2，优先采用错误概念导向策略；若整体掌握度低于0.5且处于学习初期，采用覆盖策略；若存在多日未练习的已掌握知识单元，采用间隔重复策略；若学习者请求挑战或评级多为容易，采用难度匹配策略；若需最大化信息增益诊断知识边界，采用不确定性最大化策略。',
        indent=False)
    
    # 权利要求5
    add_paragraph_zh(doc, 
        '5. 根据权利要求4所述的系统，其特征在于，所述不确定性最大化策略的计算公式为score = 1.0 - min(|P(Know) - 0.5|)，选择掌握度最接近0.5的题目以最大化信息增益。',
        indent=False)
    
    # 权利要求6
    add_paragraph_zh(doc, 
        '6. 根据权利要求1所述的系统，其特征在于，所述完整证据链追踪模块的十种事件类型包括：教学事件、知识状态更新、学习者对话、任务选择、TA尝试、评估结果、掌握度更新、错误概念激活、纠正事件、再学习事件；每个事件记录事件ID、事件类型、时间戳、关联事件ID列表、前置知识状态快照、后置知识状态快照。',
        indent=False)
    
    # 权利要求7
    add_paragraph_zh(doc, 
        '7. 根据权利要求1所述的系统，其特征在于，所述SM-2与BKT融合调度模块的掌握度调整规则为：当BKT掌握度低于0.7时，将SM-2计算的复习间隔缩短至70%；当掌握度高于0.9且评级为容易时，将复习间隔延长至120%。',
        indent=False)
    
    # 权利要求8
    add_paragraph_zh(doc, 
        '8. 根据权利要求1所述的系统，其特征在于，所述Reflect-Respond四阶段约束响应模块的约束响应生成阶段采用严格边界提示词："Strict: You may ONLY use the following knowledge in your response. Do not add external knowledge."，后跟已检索的相关事实、代码示例、已掌握知识单元列表、活跃错误概念列表。',
        indent=False)
    
    # 权利要求9
    add_paragraph_zh(doc, 
        '9. 根据权利要求1所述的系统，其特征在于，还包括智能模式切换模块，用于在教学对话模式与苏格拉底提问模式间智能切换，每累计3轮对话自动切换至提问模式一次，使用启发式模板或LLM生成Why/How类型的苏格拉底式追问。',
        indent=False)
    
    # 权利要求10
    add_paragraph_zh(doc, 
        '10. 根据权利要求1所述的系统，其特征在于，还包括教学解释多维质量评估模块，用于从准确性、完整性、清晰度三个维度评估教学输入质量，计算综合质量评分，所述质量评分作为六参数BKT中习得概率计算的加权因子。',
        indent=False)
    
    # 权利要求11 - 独立权利要求（方法）
    add_paragraph_zh(doc, 
        '11. 一种基于六参数知识追踪与语法结构约束的编程教育智能体方法，其特征在于，包括以下步骤：',
        indent=False)
    add_paragraph_zh(doc, 'S1：初始化阶段，构建六参数BKT状态矩阵，六参数初始化为默认值；', indent=True)
    add_paragraph_zh(doc, 'S2：证据链记录阶段，创建教学事件记录，关联前置知识状态快照；', indent=True)
    add_paragraph_zh(doc, 'S3：教学质量评估阶段，从准确性、完整性、清晰度三个维度评估教学输入，计算综合质量评分；', indent=True)
    add_paragraph_zh(doc, 'S4：六参数BKT更新阶段，应用六参数公式更新知识单元掌握概率，检查前置依赖满足度，若检测到误导信息则触发错误概念形成转移；', indent=True)
    add_paragraph_zh(doc, 'S5：遗忘计算阶段，应用个性化遗忘公式更新掌握概率，若低于阈值则标记遗忘状态；', indent=True)
    add_paragraph_zh(doc, 'S6：模式切换判断阶段，检查对话轮数是否满足切换条件，若满足则生成苏格拉底式问题并输出；', indent=True)
    add_paragraph_zh(doc, 'S7：Reflect-Respond四阶段处理阶段，依次执行知识提取、反思存储更新、知识检索、约束响应生成；', indent=True)
    add_paragraph_zh(doc, 'S8：任务选择阶段，从五种策略中选择适用策略，选择对应编程题目；', indent=True)
    add_paragraph_zh(doc, 'S9：代码生成阶段，触发AI智能体基于当前知识状态生成解决方案代码；', indent=True)
    add_paragraph_zh(doc, 'S10：AST守卫验证阶段，解析代码为抽象语法树，遍历节点检查所需知识单元是否为已掌握集合的子集；', indent=True)
    add_paragraph_zh(doc, 'S11：代码执行阶段，在隔离沙箱环境中执行验证通过的代码；', indent=True)
    add_paragraph_zh(doc, 'S12：执行结果-知识状态联动阶段，将执行结果反向传播至六参数BKT；', indent=True)
    add_paragraph_zh(doc, 'S13：证据链补全阶段，创建评估结果事件和掌握度更新事件，记录后置知识状态快照；', indent=True)
    add_paragraph_zh(doc, 'S14：SM-2与BKT融合调度阶段，基于当前掌握度调整复习间隔；', indent=True)
    add_paragraph_zh(doc, 'S15：循环执行S2-S14，直至达到预设学习目标。', indent=True)
    
    # 分页
    doc.add_page_break()
    
    # ==================== 摘要 ====================
    abstract_title = doc.add_paragraph()
    abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = abstract_title.add_run('摘  要')
    set_chinese_font(run, 'SimHei', 18, bold=True)
    abstract_title.paragraph_format.space_after = Pt(18)
    
    abstract_text = (
        '本发明公开了一种基于六参数知识追踪与语法结构约束的编程教育智能体系统及方法。'
        '系统在传统四参数BKT基础上扩展错误概念形成概率和错误概念纠正概率形成六参数模型，'
        '实现知识单元从习得、固化、遗忘、错误形成、纠错到再学习的完整生命周期管理。'
        '通过抽象语法树级代码守卫建立语法结构到知识单元的精确映射，实现代码复杂度与知识状态的精确匹配。'
        '采用覆盖、难度匹配、间隔重复、错误概念导向、不确定性最大化五种自适应策略动态选择学习任务。'
        '建立十种事件类型的完整证据链实现学习过程的全透明可追溯。'
        '将SM-2间隔重复算法与BKT深度融合，根据掌握度动态调整复习间隔。'
        '通过Reflect-Respond四阶段管道实现知识状态对大语言模型响应的动态约束。'
    )
    add_paragraph_zh(doc, abstract_text, indent=True)
    
    # 关键词
    keywords = doc.add_paragraph()
    run = keywords.add_run('关键词：')
    set_chinese_font(run, 'SimHei', 12, bold=True)
    run = keywords.add_run('六参数贝叶斯知识追踪；抽象语法树守卫；多策略任务选择；证据链追踪；SM-2与BKT融合；语法结构约束')
    set_chinese_font(run, 'SimSun', 12)
    keywords.paragraph_format.first_line_indent = Cm(0.74)
    
    # 保存文档
    output_path = 'e:\\cs teachable agent\\发明专利申请文件-基于六参数知识追踪与语法结构约束的编程教育智能体系统-增强版.docx'
    doc.save(output_path)
    print(f'专利申请增强版文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    create_patent_application()
